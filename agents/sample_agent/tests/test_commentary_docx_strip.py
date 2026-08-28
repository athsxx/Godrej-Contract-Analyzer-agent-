"""Tests for commentary DOCX: stripping pre-existing Word comments."""

from __future__ import annotations

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

from agents.sample_agent.redline_docx import strip_preexisting_word_comments


@pytest.mark.skipif(Document is None, reason="python-docx not installed")
def test_strip_preexisting_word_comments_clears_defs_and_markers() -> None:
    doc = Document()
    p = doc.add_paragraph("Contract body text for anchoring.")
    r = p.add_run("x")
    doc.add_comment(runs=[r], text="Prior reviewer note", author="External", initials="EX")

    strip_preexisting_word_comments(doc)

    tags = {qn("w:commentReference"), qn("w:commentRangeStart"), qn("w:commentRangeEnd")}
    assert not any(el.tag in tags for el in doc.element.body.iter())

    comments_root = doc.part.part_related_by(RT.COMMENTS).element
    assert len(list(comments_root)) == 0

    p2 = doc.add_paragraph("More body.")
    r2 = p2.add_run("y")
    doc.add_comment(runs=[r2], text="New export comment", author="Clause Analysis", initials="CA")
    assert len(list(comments_root)) >= 1
