from __future__ import annotations

import re
import unicodedata
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

try:
    from docx import Document  # type: ignore
except Exception as exc:  # pragma: no cover
    Document = None  # type: ignore
    _IMPORT_ERROR = exc
else:
    _IMPORT_ERROR = None


def _strip_evidence_metadata(text: str) -> str:
    """Remove bracketed RAG/source tags so substring matching hits contract wording."""
    s = (text or "").strip()
    if not s:
        return ""
    s = re.sub(r"\[RAG[^]]{0,800}?\]", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"\[Source:\s*[^\]]{0,800}?\]", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"\[Definitions?\s*:\s*[^\]]{0,800}?\]", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"\[PAGE\s+\d+\]", " ", s, flags=re.IGNORECASE)
    s = re.sub(r"\bFILE:\s*[^\[]+\s+PAGE\s+\d+\b", " ", s, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", s).strip()


def _flatten_doc_paragraphs(doc) -> List:
    """All paragraphs in reading order: main body blocks plus every table cell paragraph.

    ``Document.paragraphs`` omits table cells; extraction (``utils_doc_loader``) merges
    table text, so redline matching must search the same surface or most clauses never match.
    """
    if Document is None or doc is None:
        return []
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    out: List = []

    def cell_paragraphs(cell: _Cell) -> List:
        acc: List = []
        for child in cell._tc.iterchildren():
            if isinstance(child, CT_P):
                acc.append(Paragraph(child, cell))
            elif isinstance(child, CT_Tbl):
                acc.extend(table_paragraphs(Table(child, cell)))
        return acc

    def table_paragraphs(table: Table) -> List:
        acc: List = []
        for row in table.rows:
            for cell in row.cells:
                acc.extend(cell_paragraphs(cell))
        return acc

    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            out.append(Paragraph(child, doc))
        elif isinstance(child, CT_Tbl):
            out.extend(table_paragraphs(Table(child, doc)))
    return out


def strip_preexisting_word_comments(doc) -> None:
    """Remove existing Word comments and in-body markers so exports only add new comments.

    Clears ``word/comments.xml`` and strips ``w:commentReference`` / ``w:commentRangeStart`` /
    ``w:commentRangeEnd`` from the main story and from header, footer, footnote, and endnote
    parts when present. Without this, commentary export keeps supplier review comments.
    """
    if Document is None or doc is None:
        return
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn

    comment_tags = {
        qn("w:commentReference"),
        qn("w:commentRangeStart"),
        qn("w:commentRangeEnd"),
    }

    def _strip_markers(root) -> None:
        if root is None:
            return
        dead = [el for el in root.iter() if el.tag in comment_tags]
        for el in dead:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    _strip_markers(doc.element.body)

    part = doc.part
    for rel in part.rels.values():
        if getattr(rel, "is_external", False):
            continue
        rt = rel.reltype
        if rt not in (RT.HEADER, RT.FOOTER, RT.FOOTNOTES, RT.ENDNOTES):
            continue
        try:
            tgt = rel.target_part
            _strip_markers(getattr(tgt, "element", None))
        except Exception:
            continue

    try:
        comments_root = part.part_related_by(RT.COMMENTS).element
    except KeyError:
        return
    for child in list(comments_root):
        comments_root.remove(child)


def _first_evidence(evidence_snippet: str) -> str:
    text = _strip_evidence_metadata(evidence_snippet or "")
    if not text:
        return ""
    # Top-N evidence is joined by " | " in current pipeline.
    parts = [_strip_evidence_metadata(p) for p in text.split(" | ") if p.strip()]
    if not parts:
        return text
    # Prefer clause body sentence over short heading-like fragments.
    for part in parts:
        if len(part) >= 45 and " " in part:
            return part
    return parts[0]


def _align_evidence_snippet_to_document(
    evidence: str,
    flat_paras: List,
    *,
    source_contract_text: str = "",
) -> str:
    """Prefer a snippet that literally appears in the export DOCX (body + tables).

    RAG and bracketed metadata often produce evidence that is *close* to contract
    wording but not a substring of any ``Paragraph.text``. Matching then fails.
    When alignment finds a long contiguous word run from a real paragraph that
    is also contained in the normalized evidence, we use that run as the search
    string (preserves original casing from the document where word lists align).
    """
    try:
        from . import config as agent_config
        if not getattr(agent_config, "ENABLE_EVIDENCE_DOC_ALIGNMENT", True):
            return evidence
    except Exception:
        pass

    ev0 = (evidence or "").strip()
    if len(ev0) < 25:
        return evidence
    ev_inner = _strip_evidence_metadata(ev0).strip() or ev0
    ev_n = _normalize_text(ev_inner)
    if len(ev_n) < 25:
        return evidence

    raw_texts = [(p.text or "").strip() for p in flat_paras if (p.text or "").strip()]
    for raw in raw_texts:
        if len(raw) < 35:
            continue
        pn = _normalize_text(raw)
        if ev_n in pn or (len(pn) >= 60 and pn in ev_n):
            return ev_inner

    ev_words = ev_n.split()
    if len(ev_words) < 6:
        return evidence

    best_len = 0
    best_span: List[str] = []

    for raw in raw_texts:
        pn = _normalize_text(raw)
        pw = pn.split()
        if len(pw) < 6:
            continue
        for i in range(len(pw)):
            max_j = min(len(pw), i + 90)
            for j in range(i + 6, max_j + 1):
                span = pw[i:j]
                phrase = " ".join(span)
                if len(phrase) < 28:
                    continue
                if phrase not in ev_n:
                    continue
                if j - i > best_len:
                    best_len = j - i
                    best_span = list(span)

    if best_len < 6 or not best_span:
        if (source_contract_text or "").strip():
            fn = _normalize_text(source_contract_text)
            if ev_n in fn:
                return ev_inner
        return evidence

    phrase = " ".join(best_span)
    for raw in raw_texts:
        ow = raw.split()
        pn = _normalize_text(raw).split()
        if len(ow) != len(pn) or len(pn) < best_len:
            continue
        for i in range(len(pn) - best_len + 1):
            if " ".join(pn[i : i + best_len]) != phrase:
                continue
            chunk = " ".join(ow[i : i + best_len]).strip()
            if len(chunk) >= 28:
                return chunk[:2500]

    return phrase[:2500]


def _normalize_for_edit(original: str, row: Dict[str, str]) -> str:
    clause_name = (row.get("clause_name") or "").lower()
    lowered = original.lower()

    # Prefer minimal in-clause edits over wholesale rewrites (so validator accepts length/overlap).
    if "governing law" in clause_name or "jurisdiction" in clause_name:
        # Align to approved combinations (e.g. India/Mumbai)
        edited = re.sub(r"\blaws?\s+of\s+new\s+york\b", "laws of India", original, flags=re.IGNORECASE)
        edited = re.sub(r"\bcourts?\s+in\s+new\s+york\b", "courts in Mumbai", edited, flags=re.IGNORECASE)
        if edited != original:
            return edited
    if "limitation of liability" in clause_name:
        edited = re.sub(r"\b\d+(?:\.\d+)?\s*%", "100%", original, count=1)
        edited = re.sub(r"\bnot excluded\b", "excluded", edited, flags=re.IGNORECASE)
        if edited != original:
            return edited
    if "force majeure" in clause_name:
        # Match various phrasings: cash-flow, economic hardship, shortage of labor, market fluctuations
        edited = re.sub(
            r"economic hardship and cash[-\s]?flow constraints are also treated as force majeure events\.?",
            "economic hardship and cash-flow constraints are excluded from force majeure events.",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        edited = re.sub(
            r"economic hardship, shortage of labor, and market fluctuations shall also constitute force majeure events\.?",
            "economic hardship, shortage of labor, and market fluctuations are excluded from force majeure events.",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        # Shortage of labor, economic hardship, cash-flow disruptions (sample contract style)
        edited = re.sub(
            r"force majeure includes shortage of labor, economic hardship, and cash[-\s]?flow disruptions\.?",
            "force majeure excludes shortage of labor, economic hardship, and cash-flow disruptions.",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        edited = re.sub(
            r"shortage of labor, economic hardship, and cash[-\s]?flow disruptions",
            "shortage of labor, economic hardship, and cash-flow disruptions are excluded",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
    if "liquidated damages" in clause_name:
        edited = original
        edited = re.sub(r"\b2\s*%\s*of\s+the\s+total\s+contract\s+value\s+per\s*week\b", "0.5% of delayed value per week", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\b1\.5\s*%\s*per\s*week\s+of\s+total\s+order\s+value\b", "0.5% per week of delayed value", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\b1(?:\.0+)?\s*%\s*per\s*week\b", "0.5% per week", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\b12\s*%\b", "5%", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\b(?:total\s+order\s+value|total\s+contract\s+value)\b", "delayed value", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\bon\s+total\s+contract\s+value\b", "on delayed value", edited, flags=re.IGNORECASE)
        edited = re.sub(r"\bwith\s+no\s+overall\s+cap\b", "with a 5% overall cap", edited, flags=re.IGNORECASE)
        edited = re.sub(
            r",?\s*and\s+may\s+be\s+applied\s+in\s+addition\s+to\s+other\s+remedies\.?",
            ".",
            edited,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
    if "quantity protection" in clause_name:
        edited = re.sub(
            r"no reimbursement (?:shall be payable|applies) (?:to supplier for excess inventory resulting from forecast reductions|for deviations beyond plus or minus 20%)\.?",
            "reimbursement applies at actuals for deviations beyond plus or minus 20%.",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        edited = re.sub(
            r"no reimbursement applies for deviations beyond plus or minus 20%",
            "reimbursement applies at actuals for deviations beyond plus or minus 20%",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
    if "inventory requirements" in clause_name:
        edited = re.sub(r"twelve\s*\(\s*12\s*\)\s*weeks", "4 weeks", original, flags=re.IGNORECASE)
        if edited != original:
            return edited
        edited = re.sub(r"\b(\d+)\s*weeks\b", "4 weeks", original, count=1, flags=re.IGNORECASE)
        if edited != original:
            return edited
    if "change orders procedure" in clause_name:
        edited = re.sub(r"email approval", "formal signed change order approval", original, flags=re.IGNORECASE)
        edited = re.sub(
            r"can be completed later",
            "must be completed before implementation",
            edited,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
    if "orders extending beyond termination" in clause_name:
        edited = re.sub(
            r"will continue unless mutually agreed otherwise",
            "will terminate unless mutually agreed otherwise",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        edited = re.sub(
            r"all existing orders continue\s+(?:for\s+\d+\s+months\s+)?(?:after\s+termination\s+)?(?:at\s+pre[-\s]?termination\s+prices\s*)?\.?",
            "all existing orders terminate upon termination unless mutually agreed otherwise.",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited
        edited = re.sub(
            r"\bcontinue\s+for\s+24\s+months\b",
            "terminate unless mutually agreed otherwise",
            original,
            flags=re.IGNORECASE,
        )
        if edited != original:
            return edited

    # Use suggested/ideal/mitigation when they pass validator. Fallback: anchored or partial replacement.
    evidence = (row.get("evidence_text") or row.get("evidence_snippet") or "").strip()
    suggested = (row.get("suggested_text") or "").strip()
    ideal = (row.get("gb_ideal_position") or "").strip()
    mitigation = (row.get("mitigation_recommendation") or "").strip()
    for candidate in [suggested, ideal, mitigation]:
        if not candidate:
            continue
        if _is_valid_edit(original, candidate, allow_minimal_replacements=True):
            return candidate
        anchored = _build_anchored_replacement(original, candidate, evidence)
        if anchored != candidate and _is_valid_edit(original, anchored, allow_minimal_replacements=True):
            return anchored
        partial = _try_partial_replacement(original, candidate, evidence)
        if partial:
            return partial
    return original


def _write_redline(paragraph, original_text: str, edited_text: str) -> None:
    original_words = (original_text or "").split()
    edited_words = (edited_text or "").split()
    matcher = SequenceMatcher(None, original_words, edited_words)

    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == "equal":
            token = " ".join(edited_words[j1:j2]).strip()
            if token:
                paragraph.add_run(token + " ")
        elif op == "delete":
            token = " ".join(original_words[i1:i2]).strip()
            if token:
                run = paragraph.add_run(token + " ")
                run.font.strike = True
        elif op == "insert":
            token = " ".join(edited_words[j1:j2]).strip()
            if token:
                run = paragraph.add_run(token + " ")
                run.font.underline = True
        elif op == "replace":
            deleted = " ".join(original_words[i1:i2]).strip()
            inserted = " ".join(edited_words[j1:j2]).strip()
            if deleted:
                run_del = paragraph.add_run(deleted + " ")
                run_del.font.strike = True
            if inserted:
                run_ins = paragraph.add_run(inserted + " ")
                run_ins.font.underline = True


def _paragraph_comment_anchor_runs(target_paragraph):
    """Runs for ``Document.add_comment``: first and last run of the paragraph.

    python-docx only uses the first and last run to place ``w:commentRangeStart`` /
    ``w:commentRangeEnd``. Passing only the final run highlights a trailing fragment
    (often a single punctuation run); first→last spans the whole paragraph.
    """
    if not target_paragraph.runs:
        target_paragraph.add_run(" ")
    runs = list(target_paragraph.runs)
    if len(runs) == 1:
        return runs
    return [runs[0], runs[-1]]


def _add_margin_comments(doc, target_paragraph, policy_comment: str, counterfactual_comment: str) -> bool:
    if not hasattr(doc, "add_comment"):
        return False
    try:
        anchor_runs = _paragraph_comment_anchor_runs(target_paragraph)
        doc.add_comment(
            runs=anchor_runs,
            text=f"Deviation from company policy: {policy_comment}",
            author="Legal Review Agent",
            initials="LRA",
        )
        doc.add_comment(
            runs=anchor_runs,
            text=f"Counterfactual risk: {counterfactual_comment}",
            author="Legal Review Agent",
            initials="LRA",
        )
        return True
    except Exception:
        return False


def _normalize_text(text: str) -> str:
    """Lowercase, Unicode-normalize, and collapse whitespace for robust matching."""
    s = (text or "").strip().lower()
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u00a0", " ").replace("\u2009", " ").replace("\u2011", "-")
    return re.sub(r"\s+", " ", s).strip()


_TOKEN_RE = re.compile(r"[\w%]+", re.UNICODE)


def _text_tokens(text: str) -> Set[str]:
    """Word-like tokens including accented letters (French, etc.).

    The legacy pattern ``[a-z0-9%]+`` dropped non-ASCII letters, so French (and
    many legal DOCX bodies) produced empty token sets → overlap scoring and
    validators always failed and paragraph matching scored 0.
    """
    n = _normalize_text(text)
    out: Set[str] = set()
    for t in _TOKEN_RE.findall(n):
        if len(t) >= 2 or t.isdigit():
            out.add(t)
    return out


def _token_overlap_score(a: str, b: str) -> float:
    a_tokens = _text_tokens(a)
    b_tokens = _text_tokens(b)
    if not a_tokens or not b_tokens:
        return 0.0
    return len(a_tokens & b_tokens) / max(1, len(a_tokens))


def _is_heading_like(text: str) -> bool:
    clean = " ".join((text or "").split()).strip()
    if not clean:
        return True
    # Numbered headings like "1. Clause Name"
    if re.match(r"^\d+[\.\)]\s+[A-Za-z].{0,100}$", clean):
        return True
    # Very short title-like lines without sentence punctuation.
    if len(clean) <= 60 and not re.search(r"[.;:]", clean):
        alpha_ratio = sum(1 for c in clean if c.isalpha()) / max(1, len(clean))
        if alpha_ratio >= 0.6:
            return True
    return False


def _is_clause_body_paragraph(text: str) -> bool:
    clean = " ".join((text or "").split()).strip()
    if not clean:
        return False
    if _is_heading_like(clean):
        return False
    # Clause body should be sentence-like and reasonably long.
    if len(clean) < 45:
        return False
    # Some civil-law / schedule blocks run long without '.' in python-docx cell text.
    # Long comma-separated clauses are still clause-like (avoid matching title-only lines).
    has_clause_punct = bool(re.search(r"[.;:]", clean)) or (len(clean) >= 70 and "," in clean)
    if not has_clause_punct and len(clean) < 120:
        return False
    return True


def _is_clause_body_paragraph_relaxed(text: str) -> bool:
    """Like `_is_clause_body_paragraph` but allows shorter / semicolon-heavy aerospace boilerplate.

    Used only in late fallback passes so TOC rows and one-line headings stay excluded.
    """
    clean = " ".join((text or "").split()).strip()
    if not clean or _is_heading_like(clean):
        return False
    if len(clean) < 38:
        return False
    # Long uninterrupted legal blocks (common in Safran-style agreements).
    if len(clean) >= 88:
        return True
    has_clause_punct = bool(re.search(r"[.;:]", clean)) or (len(clean) >= 52 and "," in clean)
    if not has_clause_punct and len(clean) < 110:
        return False
    return True


_CLAUSE_KEYWORD_STOP = frozenset(
    {
        "this",
        "that",
        "with",
        "from",
        "shall",
        "will",
        "such",
        "hereof",
        "thereof",
        "party",
        "parties",
        "each",
        "any",
        "all",
        "per",
        "the",
        "and",
        "for",
        "are",
        "may",
        "not",
    }
)


def _clause_name_keywords(clause_name: str) -> List[str]:
    """Extract meaningful tokens from an Agent-1 clause title for proximity anchoring."""
    raw = (clause_name or "").strip()
    if not raw:
        return []
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]{4,}", raw)
    out: List[str] = []
    seen: Set[str] = set()
    for w in words:
        lw = w.lower()
        if lw in _CLAUSE_KEYWORD_STOP or lw in seen:
            continue
        seen.add(lw)
        out.append(lw)
    return out[:14]


def _clause_keyword_coverage(paragraph_text: str, keywords: List[str]) -> float:
    if not keywords:
        return 0.0
    pn = _normalize_text(paragraph_text)
    hits = sum(1 for k in keywords if _normalize_text(k) in pn)
    return hits / len(keywords)


def _evidence_matches_paragraph_loose(search_text: str, para_norm: str) -> bool:
    """Fallback when strict phrase match fails (OCR/Word runs, French/English mix)."""
    if _evidence_matches_paragraph(search_text, para_norm):
        return True
    st = _normalize_text(search_text)
    if not st or len(st) < 18:
        return True
    words = st.split()
    if len(words) < 4:
        return st in para_norm
    for i in range(0, max(1, len(words) - 2)):
        tri = " ".join(words[i : i + 3])
        if len(tri) >= 10 and tri in para_norm:
            return True
    for i in range(0, max(1, len(words) - 4)):
        pent = " ".join(words[i : i + 5])
        if len(pent) >= 14 and pent in para_norm:
            return True
    return False


def _comment_prefix_by_risk(risk_level: str) -> str:
    token = (risk_level or "").strip().lower()
    if token == "red":
        return "High-risk deviation from company policy"
    if token == "amber":
        return "Policy deviation requiring legal alignment"
    if token == "green":
        return "Aligned with company policy"
    return "Legal review note"


def _build_full_comment_text(
    *,
    risk_level: str,
    policy_comment: str,
    counterfactual_comment: str,
    mitigation: str = "",
    approval_path: str = "",
) -> str:
    """Build one concise comment: counterfactual (what-if) + approval. Minimal text for reviewer."""
    parts = []
    cf = (counterfactual_comment or "").strip()
    if cf:
        parts.append(cf)
    ap = (approval_path or "").strip()
    if ap:
        parts.append(f"Approval: {ap}")
    return "\n\n".join(parts) if parts else (policy_comment or "Deviation requires legal review.")


def _add_risk_aware_margin_comments(
    doc,
    target_paragraph,
    *,
    risk_level: str,
    policy_comment: str,
    counterfactual_comment: str,
    mitigation: str = "",
    approval_path: str = "",
    original_text: str = "",
    edited_text: str = "",
    clause_name: str = "",
    gb_ideal: str = "",
) -> bool:
    """Add one structured Word comment with policy, mitigation, counterfactual, and approval path.

    When ENABLE_SEMANTIC_COMMENTS is True, uses generate_comment_for_edit to produce a
    context-aware comment. On failure or empty result, falls back to _build_full_comment_text.
    """
    if not hasattr(doc, "add_comment"):
        return False
    try:
        anchor_runs = _paragraph_comment_anchor_runs(target_paragraph)
        text = ""
        try:
            from . import config as agent_config
            if getattr(agent_config, "ENABLE_SEMANTIC_COMMENTS", False):
                from .semantic_edit_generator import generate_comment_for_edit
                text = generate_comment_for_edit(
                    original_text=original_text,
                    edited_text=edited_text,
                    clause_name=clause_name,
                    risk_level=risk_level,
                    gb_ideal=gb_ideal,
                )
        except Exception:
            pass
        if not text:
            text = _build_full_comment_text(
                risk_level=risk_level,
                policy_comment=policy_comment,
                counterfactual_comment=counterfactual_comment,
                mitigation=mitigation,
                approval_path=approval_path,
            )
        doc.add_comment(
            runs=anchor_runs,
            text=text,
            author="Legal Review Agent",
            initials="LRA",
        )
        return True
    except Exception:
        return False


def _evidence_matches_paragraph(search_text: str, para_norm: str) -> bool:
    """Require a significant phrase from evidence to appear in the paragraph.

    Prevents wrong-clause matches (e.g. '15 days of receipt' from dispute clause
    matching LD; 'Phase 1 Definitions' matching liability evidence).
    """
    st = _normalize_text(search_text)
    if not st or len(st) < 12:
        return True  # Too short to check
    words = st.split()
    phrases: List[str] = []
    if len(words) >= 5:
        phrases.append(" ".join(words[:5]))
    else:
        phrases.append(st[:25] if len(st) >= 25 else st)
    if len(words) >= 10:
        phrases.append(" ".join(words[2:7]))
        phrases.append(" ".join(words[4:9]))
    return any(p in para_norm for p in phrases if p)


def _is_generic_evidence_for_clause(evidence: str, clause_name: str) -> bool:
    """Exclude evidence that commonly appears in other clauses (cross-clause pollution)."""
    ev = (evidence or "").strip().lower()
    cn = (clause_name or "").strip().lower()
    if not ev:
        return False
    # LD evidence "15 days" often from dispute/escalation clauses; indemnification is not LD
    if "liquidated damages" in cn or "ld" in cn:
        if "15 days" in ev or "within 15 days" in ev or "15 days of receipt" in ev or "days of receipt of the other" in ev:
            return True
        if "defend, hold harmless" in ev or "hold harmless and indemnify" in ev or ("indemnify" in ev and "liquidated" not in ev and "delay" not in ev):
            return True
        if "export license" in ev and ("withdrawn" in ev or "not renewed" in ev) and "delay" not in ev and "per week" not in ev:
            return True
    # Limitation of Liability: "loss of business which is incapable" is from consequential carve-out; LD/penalty is not LoL
    if "limitation of liability" in cn or "lod" in cn:
        if "loss of business which is incapable" in ev or "incapable of accurate estimation" in ev:
            return True
        if "per overdue calendar day" in ev or "penalty being capped at" in ev or "penalty, which does not constitute" in ev:
            return True
    # Force majeure: "180 days prior" is from design/process change clauses
    if "force majeure" in cn:
        if "180 days prior" in ev and "force majeure" not in ev and "fm" not in ev and "impediment" not in ev:
            return True
    # Inventory: manufacturing/subcontracting clauses often share "weeks" or "forecast"
    if "inventory" in cn:
        if ("manufacturing location" in ev or "subcontracting of process" in ev or "180 days prior to the proposed" in ev) and "inventory" not in ev:
            return True
    # Orders extending: "partial termination" alone is often termination procedure
    if "orders extending" in cn or "orders" in cn:
        if "partial termination" in ev and ("orders" not in ev and "in-effect" not in ev and "purchase orders" not in ev):
            return True
    # Quantity protection: "not eligible for equitable adjustment" is change-order clause
    if "quantity protection" in cn or "quantity" in cn:
        if "not eligible for equitable adjustment" in ev and "forecast" not in ev and "deviation" not in ev:
            return True
    return False


def _find_best_matching_paragraph_index(
    doc,
    evidence_text: str,
    used_indexes: Set[int],
    uploaded_position: str = "",
    clause_name: str = "",
    *,
    flat_paragraphs: Optional[List] = None,
) -> Optional[int]:
    paras = flat_paragraphs if flat_paragraphs is not None else list(doc.paragraphs)

    def _try_one(
        search_text: str,
        *,
        apply_generic_filter: bool,
        min_score: float,
        body_pred,
        loose_evidence: bool = False,
    ) -> Optional[int]:
        st = (search_text or "").strip()
        if not st:
            return None
        if apply_generic_filter and _is_generic_evidence_for_clause(st, clause_name):
            return None
        target = _normalize_text(st)
        if not target:
            return None
        best_idx: Optional[int] = None
        best_score = 0.0
        for idx, para in enumerate(paras):
            if idx in used_indexes:
                continue
            para_text = (para.text or "").strip()
            if not para_text:
                continue
            if not body_pred(para_text):
                continue
            para_norm = _normalize_text(para_text)
            if loose_evidence:
                if not _evidence_matches_paragraph_loose(st, para_norm):
                    continue
            else:
                if not _evidence_matches_paragraph(st, para_norm):
                    continue
            if target in para_norm:
                score = 1.0
            else:
                score = _token_overlap_score(target, para_norm)
            if score > best_score:
                best_idx = idx
                best_score = score
        if best_score < min_score:
            return None
        return best_idx

    def _try_clause_keyword_anchor() -> Optional[int]:
        """When evidence is filtered or fuzzy, map clause title keywords to a body paragraph."""
        kws = _clause_name_keywords(clause_name)
        if len(kws) < 2:
            return None
        ev_raw = (evidence_text or "").strip()
        up_raw = (uploaded_position or "").strip()
        best_idx: Optional[int] = None
        best_combined = 0.0
        for idx, para in enumerate(paras):
            if idx in used_indexes:
                continue
            para_text = (para.text or "").strip()
            if not para_text or not _is_clause_body_paragraph_relaxed(para_text):
                continue
            cov = _clause_keyword_coverage(para_text, kws)
            if cov < 0.14:
                continue
            o_ev = _token_overlap_score(ev_raw, para_text) if ev_raw else 0.0
            o_up = _token_overlap_score(up_raw, para_text) if up_raw and len(up_raw) > 25 else 0.0
            ov = max(o_ev, o_up)
            if ov < 0.08 and cov < 0.28:
                continue
            if not _evidence_matches_paragraph_loose(ev_raw or up_raw, _normalize_text(para_text)):
                if cov < 0.33:
                    continue
            combined = 0.48 * ov + 0.52 * cov
            if combined > best_combined:
                best_combined = combined
                best_idx = idx
        min_combined = 0.19 if len(kws) >= 4 else 0.21
        if best_idx is not None and best_combined >= min_combined:
            return best_idx
        return None

    ev = (evidence_text or "").strip()
    up = (uploaded_position or "").strip()

    idx = _try_one(ev, apply_generic_filter=True, min_score=0.32, body_pred=_is_clause_body_paragraph)
    if idx is None and up and up != ev and len(up) >= 60:
        idx = _try_one(up, apply_generic_filter=False, min_score=0.26, body_pred=_is_clause_body_paragraph)
    if idx is None and ev:
        idx = _try_one(ev, apply_generic_filter=True, min_score=0.26, body_pred=_is_clause_body_paragraph)
    # Bypass cross-clause generic filter (same paragraph must still pass overlap thresholds).
    if idx is None and ev:
        idx = _try_one(ev, apply_generic_filter=False, min_score=0.24, body_pred=_is_clause_body_paragraph)
    if idx is None and up and len(up) >= 40:
        idx = _try_one(up, apply_generic_filter=False, min_score=0.22, body_pred=_is_clause_body_paragraph)
    if idx is None and ev:
        idx = _try_one(
            ev,
            apply_generic_filter=False,
            min_score=0.20,
            body_pred=_is_clause_body_paragraph,
            loose_evidence=True,
        )
    if idx is None and up and len(up) >= 35:
        idx = _try_one(
            up,
            apply_generic_filter=False,
            min_score=0.18,
            body_pred=_is_clause_body_paragraph_relaxed,
            loose_evidence=True,
        )
    if idx is None and ev:
        idx = _try_one(
            ev,
            apply_generic_filter=False,
            min_score=0.17,
            body_pred=_is_clause_body_paragraph_relaxed,
            loose_evidence=True,
        )
    if idx is None:
        idx = _try_clause_keyword_anchor()
    return idx


def _find_best_matching_paragraph_index_with_score(
    doc,
    evidence_text: str,
    used_indexes: Set[int],
    uploaded_position: str = "",
    clause_name: str = "",
    *,
    flat_paragraphs: Optional[List] = None,
) -> Optional[Tuple[int, float]]:
    """Like _find_best_matching_paragraph_index but returns (para_idx, score) for parallel phase scoring."""
    paras = flat_paragraphs if flat_paragraphs is not None else list(doc.paragraphs)
    result = _find_best_matching_paragraph_index(
        doc,
        evidence_text,
        used_indexes,
        uploaded_position,
        clause_name,
        flat_paragraphs=flat_paragraphs,
    )
    if result is None:
        return None
    ev_raw = (evidence_text or "").strip()
    up_raw = (uploaded_position or "").strip()
    para = paras[result]
    para_text = (para.text or "").strip()
    para_norm = _normalize_text(para_text)
    ev_n = _normalize_text(ev_raw)
    up_n = _normalize_text(up_raw)
    if up_raw and len(up_raw) >= 60 and up_n in para_norm and ev_n not in para_norm:
        search_text = up_raw
    else:
        search_text = ev_raw or up_raw
    target = _normalize_text(search_text)
    if not target:
        return None
    if target in para_norm:
        score = 1.0
    else:
        score = _token_overlap_score(target, para_norm)
    return (result, score)


def _clear_paragraph_runs(paragraph) -> None:
    runs = list(paragraph.runs)
    for run in runs:
        paragraph._element.remove(run._r)


def _is_valid_edit(
    original_text: str,
    edited_text: str,
    *,
    allow_minimal_replacements: bool = False,
) -> bool:
    """Validate that an edit is clause-local and does not wholesale-replace a paragraph.

    Guardrails (documented in DOCX_REDLINE_VALIDATOR_ANALYSIS.md):
    - Original must be a clause body (length ≥ 45, sentence-like).
    - Edited must retain sufficient overlap with original (block unrelated replacements).
    - Length ratio constrains how much the edit can shrink/expand the text.
    - Min edited tokens blocks very short fragments.

    When allow_minimal_replacements=True (for suggested_text / gb_ideal_position /
    mitigation_recommendation), relaxed thresholds allow short policy lines that share
    key terms with the original while still blocking unsafe full-paragraph replacements.
    """
    if not (original_text or "").strip():
        return False
    if not (edited_text or "").strip():
        return False
    if _normalize_text(original_text) == _normalize_text(edited_text):
        return False
    if not _is_clause_body_paragraph(original_text):
        return False
    original_tokens = list(_text_tokens(original_text))
    edited_tokens = list(_text_tokens(edited_text))
    if len(original_tokens) == 0:
        return False

    # Thresholds: strict (rule-based edits) vs relaxed (policy lines from suggested/ideal)
    if allow_minimal_replacements:
        min_edited_tokens = 4
        length_ratio_min, length_ratio_max = 0.15, 3.5
        min_overlap = 0.10
        # Safety net: block very short edits with little shared content
        short_ratio_threshold, short_overlap_threshold = 0.3, 0.15
    else:
        min_edited_tokens = 8
        length_ratio_min, length_ratio_max = 0.55, 1.9
        min_overlap = 0.3
        short_ratio_threshold = short_overlap_threshold = None  # no extra guard

    if len(edited_tokens) < min_edited_tokens:
        return False
    length_ratio = len(edited_tokens) / max(1, len(original_tokens))
    if length_ratio < length_ratio_min or length_ratio > length_ratio_max:
        return False
    overlap = _token_overlap_score(original_text, edited_text)
    if short_ratio_threshold is not None and length_ratio < short_ratio_threshold:
        if overlap < short_overlap_threshold:
            return False
    return overlap >= min_overlap


def _is_valid_semantic_edit(
    original_text: str,
    edited_text: str,
    clause_name: str = "",
) -> bool:
    """Validate LLM-produced semantic edits with relaxed thresholds.

    Uses relaxed thresholds (min_overlap=0.15, length_ratio 0.2-3.0, min_edited_tokens=5)
    to allow more creative LLM edits while still blocking empty/identical/wholesale replacements.
    Keeps _is_clause_body_paragraph check for original only.
    """
    if not (original_text or "").strip():
        return False
    if not (edited_text or "").strip():
        return False
    if _normalize_text(original_text) == _normalize_text(edited_text):
        return False
    if not _is_clause_body_paragraph(original_text):
        return False
    original_tokens = list(_text_tokens(original_text))
    edited_tokens = list(_text_tokens(edited_text))
    if len(original_tokens) == 0:
        return False

    min_overlap = 0.15
    length_ratio_min, length_ratio_max = 0.2, 3.0
    min_edited_tokens = 5

    if len(edited_tokens) < min_edited_tokens:
        return False
    length_ratio = len(edited_tokens) / max(1, len(original_tokens))
    if length_ratio < length_ratio_min or length_ratio > length_ratio_max:
        return False
    overlap = _token_overlap_score(original_text, edited_text)
    return overlap >= min_overlap


def _split_sentences(text: str) -> List[str]:
    """Split text into sentence-like chunks."""
    if not (text or "").strip():
        return []
    blocks = re.split(r"[\r\n]+", text or "")
    out: List[str] = []
    for block in blocks:
        block = " ".join(block.split()).strip()
        if not block:
            continue
        for sent in re.split(r"(?<=[.!?])\s+", block):
            s = sent.strip()
            if s and len(s) >= 20:
                out.append(s)
    return out if out else ([text.strip()] if (text or "").strip() else [])


def _try_partial_replacement(
    original: str,
    replacement_text: str,
    evidence_text: str,
) -> Optional[str]:
    """Replace only the best-matching sentence when full replace fails validator."""
    if not replacement_text or not original:
        return None
    sents = _split_sentences(original)
    if len(sents) <= 1:
        return None
    evidence_norm = _normalize_text(evidence_text)
    best_idx = -1
    best_score = 0.0
    for i, sent in enumerate(sents):
        score = _token_overlap_score(evidence_norm, _normalize_text(sent))
        if score > best_score:
            best_score = score
            best_idx = i
    if best_idx < 0 or best_score < 0.1:
        return None
    new_sents = sents[:best_idx] + [replacement_text] + sents[best_idx + 1 :]
    edited = " ".join(new_sents)
    if _is_valid_edit(original, edited, allow_minimal_replacements=True):
        return edited
    return None


def _build_anchored_replacement(original: str, ideal: str, evidence_text: str) -> str:
    """Build replacement that keeps anchor phrase from original to improve overlap."""
    orig_tokens = _text_tokens(original)
    ideal_tokens = _text_tokens(ideal)
    shared = orig_tokens & ideal_tokens
    orig_words = (original or "").split()
    if len(shared) < 5 and len(orig_words) >= 8:
        anchor = " ".join(orig_words[:8])
        candidate = f"{anchor}. {ideal}"
        if _token_overlap_score(original, candidate) >= 0.2:
            return candidate
    return ideal


def _normalize_instruction(row: Dict[str, str], idx: int) -> Dict[str, str]:
    evidence_text = (
        (row.get("evidence_text") or row.get("evidence_quote") or row.get("evidence_snippet") or "").strip()
    )
    original_text = (row.get("original_text") or "").strip()
    if not original_text and evidence_text:
        original_text = evidence_text.split(" | ")[0].strip()
    uploaded = (row.get("uploaded_position") or "").strip()
    if uploaded and len(uploaded) > len(original_text) and _is_clause_body_paragraph(uploaded):
        original_text = uploaded
    return {
        "clause_id": str(row.get("clause_id") or idx),
        "clause_name": (row.get("clause_name") or f"Clause {idx}").strip(),
        "risk_level": (row.get("risk_level") or "Amber").strip(),
        "detected": (row.get("detected") or "Unclear").strip(),
        "evidence_text": evidence_text,
        "original_text": original_text,
        "suggested_text": (row.get("suggested_text") or row.get("gb_ideal_position") or "").strip(),
        "gb_ideal_position": (row.get("gb_ideal_position") or "").strip(),
        "mitigation_recommendation": (row.get("mitigation_recommendation") or "").strip(),
        "risk_rationale": (row.get("risk_rationale") or row.get("reason") or "").strip(),
        "reason": (row.get("reason") or row.get("risk_trigger") or row.get("risk_rationale") or "").strip(),
        "counterfactual": (row.get("counterfactual") or "").strip(),
        "brief_counterfactual": (row.get("brief_counterfactual") or "").strip(),
        "approval_path": (row.get("approval_path") or "").strip(),
        "confidence_score": str(row.get("confidence_score") or "").strip(),
        "uploaded_position": (row.get("uploaded_position") or "").strip(),
    }


def _instructions_from_clause_table(clause_table: List[Dict[str, str]]) -> List[Dict[str, str]]:
    return [_normalize_instruction(row, idx) for idx, row in enumerate(clause_table or [], start=1)]


def _matches_original_hint(target_paragraph_text: str, original_hint: str) -> bool:
    hint = _normalize_text(original_hint)
    para = _normalize_text(target_paragraph_text)
    if not hint:
        return True
    if hint in para:
        return True
    need = 0.34 if len(hint) > 200 else 0.36
    if _token_overlap_score(hint, para) >= need:
        return True
    return _evidence_matches_paragraph_loose(original_hint, para)


def _build_doc_from_extracted_text(text: str):
    """Build a Document with paragraph structure from extracted text (e.g. from PDF).
    Splits by double newlines, [PAGE N] blocks, or sentence boundaries so redline
    matching can find the right paragraph.
    """
    doc = Document()
    if not (text or "").strip():
        return doc
    # Split by paragraph-like boundaries: double newline, or [PAGE N] markers
    blocks = re.split(r"\n\s*\n|(?=\[PAGE\s+\d+\])", text)
    for block in blocks:
        block = (block or "").strip()
        if not block:
            continue
        # Further split long blocks (>500 chars) by sentence boundaries
        if len(block) > 500:
            sents = re.split(r"(?<=[.!?])\s+", block)
            buf: List[str] = []
            buf_len = 0
            for s in sents:
                s = s.strip()
                if not s:
                    continue
                if buf_len + len(s) > 400 and buf:
                    doc.add_paragraph(" ".join(buf))
                    buf, buf_len = [], 0
                buf.append(s)
                buf_len += len(s)
            if buf:
                doc.add_paragraph(" ".join(buf))
        else:
            doc.add_paragraph(block)
    return doc


def build_reviewed_contract_docx(
    *,
    output_path: Path,
    clause_table: List[Dict[str, str]],
    edit_instructions: Optional[List[Dict[str, str]]] = None,
    source_contract_text: str,
    source_docx_path: Optional[Path] = None,
    render_warnings: Optional[List[str]] = None,
    verification_flags: Optional[List[Dict[str, str]]] = None,
) -> Path:
    if Document is None:
        raise RuntimeError(f"python-docx is required for DOCX export: {_IMPORT_ERROR}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_docx_path and source_docx_path.exists() and source_docx_path.suffix.lower() == ".docx":
        doc = Document(str(source_docx_path))
    else:
        # Fallback for PDF/non-docx uploads: build DOCX with paragraph structure
        # from extracted text so redline matching can find clause paragraphs.
        doc = _build_doc_from_extracted_text(source_contract_text or "")

    flat_paras = _flatten_doc_paragraphs(doc) or list(doc.paragraphs)
    n_flat = len(flat_paras)

    instructions = edit_instructions or _instructions_from_clause_table(clause_table)
    used_paragraph_indexes: Set[int] = set()
    for i, raw in enumerate(instructions, start=1):
        row = _normalize_instruction(raw, idx=i)
        risk_level = row["risk_level"]
        if risk_level == "Green":
            # Green rows should not be redlined or commented.
            continue
        if row["detected"] in {"No", "Unclear"}:
            if render_warnings is not None:
                render_warnings.append(
                    f"{row['clause_name']}: skipped redline because clause was not confidently detected."
                )
            continue

        evidence = _first_evidence(row.get("evidence_text") or "")
        evidence = _align_evidence_snippet_to_document(
            evidence, flat_paras, source_contract_text=source_contract_text or ""
        )
        if not evidence or re.match(r"^No direct matching", evidence, re.IGNORECASE):
            if render_warnings is not None:
                render_warnings.append(f"{row['clause_name']}: no valid evidence snippet for redline.")
            continue

        # Semantic: validate evidence actually describes this clause (reduces cross-clause pollution)
        try:
            from . import config as agent_config
            if getattr(agent_config, "ENABLE_EVIDENCE_CLAUSE_VALIDATION", False):
                from .semantic_edit_generator import validate_evidence_for_clause
                if not validate_evidence_for_clause(evidence, row.get("clause_name", "")):
                    if render_warnings is not None:
                        render_warnings.append(
                            f"{row['clause_name']}: evidence snippet does not describe clause (semantic validation)."
                        )
                    continue
        except Exception:
            pass

        target_idx = _find_best_matching_paragraph_index(
            doc,
            evidence,
            used_paragraph_indexes,
            uploaded_position=(row.get("uploaded_position") or "").strip(),
            clause_name=(row.get("clause_name") or "").strip(),
            flat_paragraphs=flat_paras,
        )
        if target_idx is None:
            if render_warnings is not None:
                render_warnings.append(f"{row['clause_name']}: no matching clause body paragraph found.")
            continue

        target_para = flat_paras[target_idx]
        original_text = (target_para.text or "").strip() or evidence
        if not _matches_original_hint(original_text, row["original_text"]):
            if render_warnings is not None:
                render_warnings.append(
                    f"{row['clause_name']}: matched paragraph did not satisfy original-text validator."
                )
            continue
        # Semantic-first edit flow: try semantic (LLM) first when enabled, else rule-based
        edited_text = ""
        used_semantic = False
        try:
            from . import config as agent_config
            if getattr(agent_config, "ENABLE_SEMANTIC_EDIT_GENERATION", False):
                edit_strategy = getattr(agent_config, "EDIT_STRATEGY", "semantic_first")
                if edit_strategy == "semantic_first":
                    # Try semantic edit FIRST
                    from .semantic_edit_generator import generate_semantic_edit
                    surrounding_parts: List[str] = []
                    if target_idx > 0:
                        prev = (flat_paras[target_idx - 1].text or "").strip()
                        if prev:
                            surrounding_parts.append(prev)
                    if target_idx + 1 < n_flat:
                        nxt = (flat_paras[target_idx + 1].text or "").strip()
                        if nxt:
                            surrounding_parts.append(nxt)
                    surrounding_context = "\n".join(surrounding_parts)
                    semantic_edited = generate_semantic_edit(
                        original_text=original_text,
                        clause_name=row.get("clause_name", ""),
                        gb_ideal=row.get("gb_ideal_position", "") or row.get("suggested_text", ""),
                        surrounding_context=surrounding_context,
                        risk_rationale=row.get("risk_rationale", "") or row.get("reason", ""),
                        mitigation_recommendation=row.get("mitigation_recommendation", ""),
                    )
                    if semantic_edited and _is_valid_semantic_edit(
                        original_text, semantic_edited, row.get("clause_name", "")
                    ):
                        edited_text = semantic_edited
                        used_semantic = True
                if not edited_text or not used_semantic:
                    # Fallback: rule-based
                    edited_text = _normalize_for_edit(original_text, row)
                    if edit_strategy == "rule_first":
                        # Rule-first: if rule-based produced no change, try semantic
                        orig_norm = re.sub(r"\s+", " ", (original_text or "").strip())
                        edit_norm = re.sub(r"\s+", " ", (edited_text or "").strip())
                        if not edit_norm or edit_norm == orig_norm:
                            from .semantic_edit_generator import generate_semantic_edit
                            surrounding_parts = []
                            if target_idx > 0:
                                prev = (flat_paras[target_idx - 1].text or "").strip()
                                if prev:
                                    surrounding_parts.append(prev)
                            if target_idx + 1 < n_flat:
                                nxt = (flat_paras[target_idx + 1].text or "").strip()
                                if nxt:
                                    surrounding_parts.append(nxt)
                            surrounding_context = "\n".join(surrounding_parts)
                            semantic_edited = generate_semantic_edit(
                                original_text=original_text,
                                clause_name=row.get("clause_name", ""),
                                gb_ideal=row.get("gb_ideal_position", "") or row.get("suggested_text", ""),
                                surrounding_context=surrounding_context,
                                risk_rationale=row.get("risk_rationale", "") or row.get("reason", ""),
                                mitigation_recommendation=row.get("mitigation_recommendation", ""),
                            )
                            if semantic_edited and _is_valid_semantic_edit(
                                original_text, semantic_edited, row.get("clause_name", "")
                            ):
                                edited_text = semantic_edited
                                used_semantic = True
            if not edited_text:
                edited_text = _normalize_for_edit(original_text, row)
        except Exception:
            if not edited_text:
                edited_text = _normalize_for_edit(original_text, row)

        validator_passes = (
            _is_valid_semantic_edit(original_text, edited_text, row.get("clause_name", ""))
            if used_semantic
            else _is_valid_edit(original_text, edited_text, allow_minimal_replacements=True)
        )
        try:
            from . import config as agent_config
            if (
                getattr(agent_config, "ENABLE_SEMANTIC_FALLBACK_TO_RULE", True)
                and used_semantic
                and not validator_passes
            ):
                rule_fb = _normalize_for_edit(original_text, row)
                if _normalize_text(rule_fb) != _normalize_text(original_text) and _is_valid_edit(
                    original_text, rule_fb, allow_minimal_replacements=True
                ):
                    edited_text = rule_fb
                    used_semantic = False
                    validator_passes = _is_valid_edit(
                        original_text, edited_text, allow_minimal_replacements=True
                    )
                    if render_warnings is not None:
                        render_warnings.append(
                            f"{row['clause_name']}: semantic edit failed validation; applied rule-based redline instead."
                        )
        except Exception:
            pass
        # Bypass: rule-based minimal edits (reasonable overlap) are trusted
        if not validator_passes and not used_semantic and edited_text and original_text:
            overlap = _token_overlap_score(original_text, edited_text)
            if overlap >= 0.25 and len(edited_text) >= 15 and edited_text.strip() != original_text.strip():
                validator_passes = True
        # Always reserve this paragraph to avoid double-comment/redline from later clauses.
        used_paragraph_indexes.add(target_idx)

        agent4_flagged = False
        if validator_passes and verification_flags is not None:
            try:
                from . import config as agent_config
                if getattr(agent_config, "ENABLE_AGENT4_VERIFICATION", False):
                    from .agent4_verifier import verify_redline_edit
                    surrounding_parts = []
                    if target_idx > 0:
                        prev = (flat_paras[target_idx - 1].text or "").strip()
                        if prev:
                            surrounding_parts.append(f"[Before] {prev}")
                    if target_idx + 1 < n_flat:
                        nxt = (flat_paras[target_idx + 1].text or "").strip()
                        if nxt:
                            surrounding_parts.append(f"[After] {nxt}")
                    surrounding_context = "\n".join(surrounding_parts) if surrounding_parts else "(no surrounding context)"
                    result = verify_redline_edit(
                        original_text=original_text,
                        edited_text=edited_text,
                        clause_name=row.get("clause_name", ""),
                        surrounding_context=surrounding_context,
                        gb_ideal=row.get("gb_ideal_position", "") or row.get("suggested_text", ""),
                    )
                    if result.get("verdict") == "flag":
                        agent4_flagged = getattr(agent_config, "SKIP_REDLINE_WHEN_AGENT4_FLAGS", True)
                        verification_flags.append({
                            "clause_name": row.get("clause_name", ""),
                            "reason": result.get("reason", ""),
                            "suggestion": result.get("suggestion", ""),
                        })
            except Exception:
                pass

        # Optional: LLM semantic validation (preserves legal intent?)
        edit_semantic_rejected = False
        if validator_passes and not agent4_flagged:
            try:
                from . import config as agent_config
                if getattr(agent_config, "ENABLE_EDIT_SEMANTIC_VALIDATION", False):
                    from .semantic_edit_generator import validate_edit_semantics
                    gb_ideal = (row.get("gb_ideal_position") or row.get("suggested_text") or "").strip()
                    if not validate_edit_semantics(
                        original_text=original_text,
                        edited_text=edited_text,
                        clause_name=row.get("clause_name", ""),
                        gb_ideal=gb_ideal,
                    ):
                        edit_semantic_rejected = True
            except Exception:
                pass

        if edit_semantic_rejected:
            try:
                from . import config as agent_config
                if getattr(agent_config, "ENABLE_SEMANTIC_FALLBACK_TO_RULE", True):
                    rule_fb = _normalize_for_edit(original_text, row)
                    if _normalize_text(rule_fb) != _normalize_text(original_text) and _is_valid_edit(
                        original_text, rule_fb, allow_minimal_replacements=True
                    ):
                        edited_text = rule_fb
                        used_semantic = False
                        edit_semantic_rejected = False
                        validator_passes = _is_valid_edit(
                            original_text, edited_text, allow_minimal_replacements=True
                        )
                        if render_warnings is not None:
                            render_warnings.append(
                                f"{row['clause_name']}: intent validator rejected LLM edit; applied rule-based redline instead."
                            )
            except Exception:
                pass

        # Apply redline only if validator passes, Agent 4 did not flag, and semantic validator passed
        if validator_passes and not agent4_flagged and not edit_semantic_rejected:
            _clear_paragraph_runs(target_para)
            _write_redline(target_para, original_text, edited_text)
        elif render_warnings is not None:
            render_warnings.append(
                f"{row['clause_name']}: validator rejected proposed edit (overlap/length/sentence guardrail)."
            )

        # Add comment: brief counterfactual only (minimal text). Use brief_counterfactual if available.
        brief_cf = (row.get("brief_counterfactual") or "").strip()
        if not brief_cf:
            raw_cf = (row.get("counterfactual") or "").strip()
            if len(raw_cf) > 100:
                first = raw_cf.split(".")[0].split(";")[0].strip()
                brief_cf = first[:100] + "..." if len(first) > 100 else first
            else:
                brief_cf = raw_cf
        counterfactual_comment = brief_cf or "Align clause to GB policy."
        has_margin_comments = _add_risk_aware_margin_comments(
            doc,
            target_para,
            risk_level=risk_level,
            policy_comment="",  # Omitted to reduce comment length
            counterfactual_comment=counterfactual_comment,
            mitigation="",  # Omitted; full mitigation in checklist/CSV
            approval_path=row.get("approval_path") or "",
            original_text=original_text,
            edited_text=edited_text,
            clause_name=row.get("clause_name") or "",
            gb_ideal=(row.get("gb_ideal_position") or row.get("suggested_text") or "").strip(),
        )
        if not has_margin_comments and render_warnings is not None:
            render_warnings.append(
                f"{row['clause_name']}: margin comments unavailable in current python-docx runtime; skipped comment insertion."
            )

    doc.save(str(output_path))
    return output_path


def build_reviewed_contract_docx_parallel(
    *,
    output_path: Path,
    clause_table: List[Dict[str, str]],
    edit_instructions: Optional[List[Dict[str, str]]] = None,
    source_contract_text: str,
    source_docx_path: Optional[Path] = None,
    render_warnings: Optional[List[str]] = None,
    verification_flags: Optional[List[Dict[str, str]]] = None,
) -> Path:
    """Parallel redline flow: Phase 1–3 via orchestrator, Phase 4 applies redlines/comments sequentially."""
    if Document is None:
        raise RuntimeError(f"python-docx is required for DOCX export: {_IMPORT_ERROR}")

    from .orchestrator import (
        run_redline_phase1_match,
        run_redline_phase2_resolve,
        run_redline_phase3_edit,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_docx_path and source_docx_path.exists() and source_docx_path.suffix.lower() == ".docx":
        doc = Document(str(source_docx_path))
    else:
        doc = _build_doc_from_extracted_text(source_contract_text or "")

    flat_paras = _flatten_doc_paragraphs(doc) or list(doc.paragraphs)

    instructions = edit_instructions or _instructions_from_clause_table(clause_table)

    # Phase 1: parallel evidence validation + paragraph matching
    phase1_results = run_redline_phase1_match(
        instructions,
        doc,
        flat_paragraphs=flat_paras,
        source_contract_text=source_contract_text or "",
    )

    # Phase 2: resolve paragraph conflicts (greedy by score, prefer higher risk)
    assignments = run_redline_phase2_resolve(phase1_results, instructions)

    # Phase 3: parallel semantic edit + Agent4 per assigned clause
    phase3_results = run_redline_phase3_edit(instructions, doc, assignments, flat_paragraphs=flat_paras)

    # Phase 4: apply redlines and comments sequentially (document order)
    phase3_results.sort(key=lambda x: x[1])  # by para_idx

    for clause_idx, para_idx, original_text, edited_text, row, agent4_flagged, agent4_result in phase3_results:
        target_para = flat_paras[para_idx]
        risk_level = row.get("risk_level", "Amber")

        if not agent4_flagged:
            _clear_paragraph_runs(target_para)
            _write_redline(target_para, original_text, edited_text)

        if verification_flags is not None and agent4_flagged:
            verification_flags.append({
                "clause_name": row.get("clause_name", ""),
                "reason": agent4_result.get("reason", ""),
                "suggestion": agent4_result.get("suggestion", ""),
            })

        brief_cf = (row.get("brief_counterfactual") or "").strip()
        if not brief_cf:
            raw_cf = (row.get("counterfactual") or "").strip()
            if len(raw_cf) > 100:
                first = raw_cf.split(".")[0].split(";")[0].strip()
                brief_cf = first[:100] + "..." if len(first) > 100 else first
            else:
                brief_cf = raw_cf
        counterfactual_comment = brief_cf or "Align clause to GB policy."
        _add_risk_aware_margin_comments(
            doc,
            target_para,
            risk_level=risk_level,
            policy_comment="",
            counterfactual_comment=counterfactual_comment,
            mitigation="",
            approval_path=row.get("approval_path") or "",
            original_text=original_text,
            edited_text=edited_text,
            clause_name=row.get("clause_name") or "",
            gb_ideal=(row.get("gb_ideal_position") or row.get("suggested_text") or "").strip(),
        )

    doc.save(str(output_path))
    return output_path


def _query_for_supporting_excerpt(row: Dict[str, str]) -> str:
    parts = [
        (row.get("clause_name") or "").strip(),
        (row.get("evidence_quote") or "").strip()[:520],
        (row.get("risk_rationale") or "").strip()[:360],
    ]
    return " ".join(p for p in parts if p).strip()


def _score_supporting_window(window: str, query: str) -> float:
    q_tokens = set(re.findall(r"[a-z0-9]{4,}", (query or "").lower()))
    if not q_tokens:
        return 0.0
    w_lower = (window or "").lower()
    hits = sum(1 for t in q_tokens if t in w_lower)
    return hits / max(1, min(len(q_tokens), 28))


def _appendix_crossref_suffix(fragment: str, missing_docs: Optional[List[dict]]) -> str:
    if not fragment or not missing_docs:
        return ""
    frag_l = fragment.lower()
    notes: List[str] = []
    for m in re.finditer(r"(?i)\bappendix\s+([0-9]+(?:bis|ter)?)\b", frag_l):
        num = m.group(1).lower()
        for md in missing_docs:
            lab = str(md.get("label") or "")
            if not lab:
                continue
            low = lab.lower()
            if num in re.sub(r"[^a-z0-9]+", " ", low) and (
                "appendix" in low or "annex" in low or "schedule" in low
            ):
                notes.append(f"Appendix {m.group(1)}: cross-reference only — not in upload set for this session.")
                break
    if not notes:
        return ""
    return " " + " ".join(dict.fromkeys(notes))


def _supporting_span_location_lines(fname: str, sample: str, span_start: int, span_end: int) -> str:
    """Human-readable position of an excerpt inside the flattened supporting-doc text slice ``sample``."""
    sample = sample or ""
    if not sample.strip():
        return ""
    span_start = max(0, min(span_start, len(sample) - 1))
    span_end = max(span_start + 1, min(span_end, len(sample)))
    blk_before = sample[:span_start].count("\n\n")
    blk_total = max(1, sample.count("\n\n") + 1)
    pg_m = list(re.finditer(r"\[PAGE\s+(\d+)\]", sample[:span_start], flags=re.IGNORECASE))
    pg = pg_m[-1].group(1) if pg_m else None
    lines = [
        f"Where in supporting file «{fname}» (plain-text extract; first {len(sample)} characters used for search):",
        f"- Approx. blank-line block index in this extract: {blk_before + 1} of {blk_total} (split on double newlines).",
        f"- Character span in this extract slice (0-based): {span_start}–{span_end}.",
    ]
    if pg:
        lines.append(f"- Nearest preceding page marker in extract: [PAGE {pg}].")
    lines.append(
        "- Point / clause numbers: not auto-detected; search this excerpt in the Word/PDF source if you need numbering."
    )
    return "\n".join(lines)


def _supporting_excerpt_for_commentary(
    row: Dict[str, str],
    supporting_doc_texts: Optional[Dict[str, str]],
    *,
    max_chars: int = 900,
    missing_docs: Optional[List[dict]] = None,
    uploaded_filenames: Optional[List[str]] = None,
) -> str:
    """Pull a short window from supporting uploads; score windows vs clause + evidence (not first keyword hit)."""
    _ = uploaded_filenames  # reserved for future filename-aware gating
    if not supporting_doc_texts:
        return "No supporting schedules were uploaded for cross-reference."
    query = _query_for_supporting_excerpt(row)
    name = (row.get("clause_name") or "").lower()
    rationale = (row.get("risk_rationale") or "").lower()
    needles: List[str] = []
    for phrase in (
        "liquidated",
        "damages",
        "liability",
        "force majeure",
        "forecast",
        "inventory",
        "change order",
        "governing law",
        "arbitration",
        "termination",
        "price",
        "firm",
    ):
        if phrase in name or phrase in rationale:
            needles.append(phrase)
    if not needles:
        needles = [w for w in re.findall(r"[a-z]{5,}", name) if w not in {"clause", "damages", "limitation"}][:2]
    if not needles:
        needles = ["schedule"]

    chunks: List[str] = []
    win = 520
    for fname, body in supporting_doc_texts.items():
        b = body or ""
        if len(b) < 40:
            continue
        cap = min(len(b), 32000)
        sample = b[:cap]
        best_score = -1.0
        best_start = 0
        step = 100
        for start in range(0, max(1, len(sample) - 60), step):
            window = sample[start : start + win]
            sc = _score_supporting_window(window, query)
            if sc > best_score:
                best_score = sc
                best_start = start
        frag: Optional[str] = None
        span_start = 0
        span_end = 0
        if best_score >= 0.06:
            span_start = best_start
            raw = sample[span_start : span_start + win]
            span_end = span_start + len(raw)
            frag = " ".join(raw.split())
        else:
            lower = sample.lower()
            best_pos = -1
            for nd in needles:
                pos = lower.find(nd)
                if pos >= 0 and (best_pos < 0 or pos < best_pos):
                    best_pos = pos
            if best_pos >= 0:
                span_start = max(0, best_pos - 140)
                span_end = min(len(sample), best_pos + 380)
                frag = " ".join(sample[span_start:span_end].split())
        if not frag:
            continue
        prefix = f"Excerpt from uploaded file: {fname}. Does not imply other appendices were uploaded unless separately listed. "
        suffix = _appendix_crossref_suffix(frag, missing_docs)
        loc = _supporting_span_location_lines(fname, sample, span_start, span_end)
        chunks.append(prefix + f"…{frag}…" + suffix + "\n" + loc)
        if sum(len(c) for c in chunks) >= max_chars:
            break
    if not chunks:
        return "Supporting files did not contain an obvious textual hook for this clause (check OCR or filename-only schedules)."
    text = "\n".join(chunks)
    return text[:max_chars] + ("…" if len(text) > max_chars else "")


def _add_single_paragraph_comment(doc, target_para, text: str) -> bool:
    if not hasattr(doc, "add_comment"):
        return False
    try:
        anchor_runs = _paragraph_comment_anchor_runs(target_para)
        doc.add_comment(
            runs=anchor_runs,
            text=(text or "")[:8800],
            author="Clause Analysis",
            initials="CA",
        )
        return True
    except Exception:
        return False


def build_contract_commentary_docx(
    *,
    output_path: Path,
    clause_table: List[Dict[str, str]],
    source_contract_text: str,
    source_docx_path: Optional[Path] = None,
    supporting_doc_texts: Optional[Dict[str, str]] = None,
    render_warnings: Optional[List[str]] = None,
    uploaded_filenames: Optional[List[str]] = None,
    commentary_export_style: Optional[str] = None,
    primary_upload_display_name: Optional[str] = None,
) -> Path:
    """Copy primary contract DOCX and attach Word comments (no body redlines).

    Comments attach to the matched clause paragraph. Export style
    ``CONTRACT_COMMENTARY_EXPORT_STYLE`` (default ``counsel_bubble``):
    ``counsel_bubble`` — deduplicated anchor + legal brief + compact analysis record;
    ``full`` / ``counsel_short`` — legacy verbose blocks with supporting excerpts for LLM context.
    """
    if Document is None:
        raise RuntimeError(f"python-docx is required for DOCX export: {_IMPORT_ERROR}")

    from .legal_commentary import (
        assemble_legal_comment_body,
        build_comment_anchor_counsel_bubble,
        build_comment_anchor_section,
        build_comment_analysis_record_counsel_bubble,
        build_comment_facts_for_llm,
        build_comment_facts_for_llm_counsel_bubble,
        build_comment_facts_section_compact,
        format_evidence_provenance_block,
        synthesize_legal_brief_llm,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_docx_path and source_docx_path.exists() and source_docx_path.suffix.lower() == ".docx":
        doc = Document(str(source_docx_path))
    else:
        doc = _build_doc_from_extracted_text(source_contract_text or "")

    strip_preexisting_word_comments(doc)

    flat_paras = _flatten_doc_paragraphs(doc) or list(doc.paragraphs)
    used_paragraph_indexes: Set[int] = set()

    missing_docs: List[dict] = []
    for raw in clause_table or []:
        md = raw.get("_missing_supporting_docs")
        if md:
            missing_docs = list(md)
            break

    try:
        from . import config as agent_config

        export_style = (commentary_export_style or "").strip().lower() or getattr(
            agent_config, "CONTRACT_COMMENTARY_EXPORT_STYLE", "counsel_bubble"
        )
    except Exception:
        export_style = (commentary_export_style or "counsel_bubble").strip().lower() or "counsel_bubble"
    if export_style not in {"full", "counsel_short", "counsel_bubble"}:
        export_style = "counsel_bubble"

    primary_label = (primary_upload_display_name or "").strip()
    if not primary_label and source_docx_path and source_docx_path.exists():
        primary_label = source_docx_path.name

    for i, raw in enumerate(clause_table or [], start=1):
        raw_d = dict(raw)
        row = _normalize_instruction(raw_d, idx=i)
        display_row = {**raw_d, **row}

        risk_level = (row.get("risk_level") or "").strip()
        if risk_level == "Green":
            continue

        det = (row.get("detected") or "").strip()
        if det in {"No", "Unclear"}:
            if render_warnings is not None:
                render_warnings.append(
                    f"{row.get('clause_name', 'Clause')}: skipped commentary export because clause was not confidently detected."
                )
            continue

        evidence = _first_evidence(row.get("evidence_text") or "")
        evidence = _align_evidence_snippet_to_document(
            evidence, flat_paras, source_contract_text=source_contract_text or ""
        )
        if not evidence or len(evidence) < 12:
            up = (row.get("uploaded_position") or "").strip()
            evidence = _align_evidence_snippet_to_document(
                up[:220], flat_paras, source_contract_text=source_contract_text or ""
            )
        if not evidence or re.match(r"^No direct matching", evidence, re.IGNORECASE):
            if render_warnings is not None:
                render_warnings.append(
                    f"{row.get('clause_name', 'Clause')}: no valid evidence snippet for commentary comment."
                )
            continue

        up_raw = (row.get("uploaded_position") or "").strip()
        ev_raw = (evidence or "").strip()
        use_up = ""
        if up_raw and len(up_raw) < 900:
            use_up = up_raw
        elif up_raw and ev_raw and _token_overlap_score(up_raw, ev_raw) >= 0.22:
            use_up = up_raw[:1400]

        try:
            from . import config as agent_config
            if getattr(agent_config, "ENABLE_EVIDENCE_CLAUSE_VALIDATION", False):
                from .semantic_edit_generator import validate_evidence_for_clause
                if not validate_evidence_for_clause(evidence, row.get("clause_name", "")):
                    if render_warnings is not None:
                        render_warnings.append(
                            f"{row.get('clause_name', 'Clause')}: evidence snippet does not describe clause (semantic validation)."
                        )
                    continue
        except Exception:
            pass

        target_idx = _find_best_matching_paragraph_index(
            doc,
            evidence,
            used_paragraph_indexes,
            uploaded_position=use_up,
            clause_name=(row.get("clause_name") or "").strip(),
            flat_paragraphs=flat_paras,
        )
        if target_idx is None:
            if render_warnings is not None:
                render_warnings.append(f"{row.get('clause_name', 'Clause')}: could not place commentary in DOCX body.")
            continue
        target_para = flat_paras[target_idx]
        para_plain = (target_para.text or "").strip() or evidence
        if not _matches_original_hint(para_plain, row.get("original_text") or ""):
            if render_warnings is not None:
                render_warnings.append(
                    f"{row.get('clause_name', 'Clause')}: matched paragraph did not satisfy original-text validator for commentary."
                )
            continue

        used_paragraph_indexes.add(target_idx)

        if export_style == "counsel_bubble":
            anchor = build_comment_anchor_counsel_bubble(
                clause_name=(row.get("clause_name") or "").strip(),
                anchor_index=target_idx,
                total_paragraphs=len(flat_paras),
                anchored_paragraph_text=para_plain,
                evidence_for_sentence_match=evidence,
                row=display_row,
                primary_upload_display_name=primary_label,
            )
            facts = build_comment_analysis_record_counsel_bubble(
                display_row,
                uploaded_filenames=uploaded_filenames,
                primary_upload_display_name=primary_label,
            )
            try:
                facts_llm = build_comment_facts_for_llm_counsel_bubble(
                    display_row,
                    evidence_aligned=evidence,
                    primary_upload_display_name=primary_label,
                    uploaded_filenames=uploaded_filenames,
                )
                brief = synthesize_legal_brief_llm(
                    clause_name=(row.get("clause_name") or "").strip(),
                    facts_for_llm=facts_llm,
                )
            except Exception:
                brief = None
            body = assemble_legal_comment_body(
                anchor_section=anchor,
                facts_section=facts,
                llm_brief=brief,
                comment_layout="counsel_bubble",
            )
        else:
            supp = _supporting_excerpt_for_commentary(
                display_row,
                supporting_doc_texts,
                missing_docs=missing_docs or None,
                uploaded_filenames=uploaded_filenames,
            )
            prov_block = format_evidence_provenance_block(
                display_row,
                primary_upload_display_name=primary_label,
                evidence_exact_text=evidence,
                primary_flat_paragraph_1based=target_idx + 1,
                primary_flat_total_paragraphs=len(flat_paras),
            )
            anchor = build_comment_anchor_section(
                clause_name=(row.get("clause_name") or "").strip(),
                anchor_index=target_idx,
                total_paragraphs=len(flat_paras),
                anchored_paragraph_text=para_plain,
                evidence_for_sentence_match=evidence,
                evidence_provenance_block=prov_block,
            )
            facts = build_comment_facts_section_compact(
                display_row,
                supp,
                evidence_aligned=evidence,
                export_style=export_style,
                primary_upload_display_name=primary_label,
                evidence_exact_for_provenance=evidence,
                primary_flat_paragraph_1based=target_idx + 1,
                primary_flat_total_paragraphs=len(flat_paras),
            )
            try:
                facts_llm = build_comment_facts_for_llm(
                    display_row,
                    supp,
                    evidence_aligned=evidence,
                    primary_upload_display_name=primary_label,
                    evidence_exact_for_provenance=evidence,
                    primary_flat_paragraph_1based=target_idx + 1,
                    primary_flat_total_paragraphs=len(flat_paras),
                )
                brief = synthesize_legal_brief_llm(
                    clause_name=(row.get("clause_name") or "").strip(),
                    facts_for_llm=facts_llm,
                )
            except Exception:
                brief = None
            body = assemble_legal_comment_body(
                anchor_section=anchor,
                facts_section=facts,
                llm_brief=brief,
            )

        if not _add_single_paragraph_comment(doc, target_para, body) and render_warnings is not None:
            render_warnings.append(
                f"{row.get('clause_name', 'Clause')}: python-docx could not insert a comment (runtime limitation)."
            )

    doc.save(str(output_path))
    return output_path
