"""Sample conversational agent scaffold (POC – local dev only).

Uses a local LLM (Ollama-compatible) and optional RAG over the POC knowledge
document. Config via env: LOCAL_LLM_BASE_URL, LOCAL_LLM_MODEL, etc. No AWS.
"""

from __future__ import annotations

import base64
import csv
import json
import logging
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence

# Optional heavy deps: keep guarded so the scaffold can import even if missing.
try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    fitz = None  # type: ignore

try:
    import textract  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    textract = None  # type: ignore

try:
    import docx2txt  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    docx2txt = None  # type: ignore

try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    DocxDocument = None  # type: ignore

# Reuse shared extractor utilities when available.
try:
    from utils_doc_loader import extract_text as shared_extract_text  # type: ignore
    from utils_doc_loader import render_pdf_images as shared_render_pdf_images  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    shared_extract_text = None  # type: ignore
    shared_render_pdf_images = None  # type: ignore

# -------------------------------
# Basic configuration
# -------------------------------
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".csv", ".xlsx"}
MAX_UPLOAD_FILES = 10
MAX_UPLOAD_BYTES = 50 * 1024 * 1024
HISTORY_LIMIT = 25

PROJECT_ROOT = Path(__file__).resolve().parents[2]
UPLOAD_ROOT = PROJECT_ROOT / "media" / "custom_dev_sample_agent"
UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)

logger = logging.getLogger(__name__)
DEBUG_AGENT = os.environ.get("DEBUG_AGENT", "0").strip().lower() in {"1", "true", "yes", "on"}
if DEBUG_AGENT:
    logging.basicConfig(level=logging.DEBUG)


def _preview(text: str | None, limit: int = 220) -> str:
    if not text:
        return ""
    compact = " ".join(str(text).split())
    if len(compact) <= limit:
        return compact
    return compact[:limit] + "..."


def _sample_text_for_context(text: str, limit: int, *, label: str = "document") -> str:
    """Sample start + middle + end when text exceeds limit, so main content is not lost.
    Ensures LLM sees beginning, middle, and end of long documents.
    """
    if not text or len(text) <= limit:
        return (text or "").strip()
    n = len(text)
    # Allocate ~35% start, ~30% middle, ~35% end
    head_size = limit * 35 // 100
    tail_size = limit * 35 // 100
    mid_size = limit - head_size - tail_size - 80  # reserve for separators
    if mid_size < 200:
        mid_size = 200
        head_size = (limit - mid_size - 80) // 2
        tail_size = limit - mid_size - 80 - head_size
    head = text[:head_size]
    tail = text[-tail_size:] if tail_size > 0 else ""
    mid_start = max(head_size, (n - mid_size) // 2)
    mid_end = min(n - tail_size, mid_start + mid_size)
    mid = text[mid_start:mid_end]
    return (
        f"[{label} - start]\n{head}\n\n"
        f"[{label} - middle]\n{mid}\n\n"
        f"[{label} - end]\n{tail}"
    )


@dataclass
class UploadedArtifact:
    name: str
    path: Path
    mime: Optional[str] = None
    size: Optional[int] = None
    text_preview: Optional[str] = None
    # Store anything else you need (e.g., embeddings, page ranges, metadata)
    extras: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SessionState:
    files: List[UploadedArtifact] = field(default_factory=list)
    history: List[Dict[str, str]] = field(default_factory=list)
    # Place to attach in-memory indexes/vectors; wire up in TODO sections.
    vector_index: Any | None = None


_SESSION_DATA: Dict[str, SessionState] = {}


def _ensure_session(session_id: str) -> SessionState:
    state = _SESSION_DATA.setdefault(session_id, SessionState())
    # Keep chat history bounded.
    if len(state.history) > HISTORY_LIMIT:
        state.history = state.history[-HISTORY_LIMIT:]
    return state


def _chat_history_disk_path(session_id: str) -> Path:
    return UPLOAD_ROOT / session_id / "_chat_history.json"


def _load_chat_history_from_disk_if_empty(session_id: str, state: SessionState) -> None:
    """Rehydrate chat after a Celery worker wrote analysis to disk (separate process)."""
    if state.history:
        return
    path = _chat_history_disk_path(session_id)
    if not path.is_file():
        return
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, list):
            state.history = data
    except Exception:
        pass


def _persist_chat_history_to_disk(session_id: str, state: SessionState) -> None:
    if not state.history:
        return
    path = _chat_history_disk_path(session_id)
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(state.history, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _persist_upload_roles_manifest(session_id: str, state: SessionState) -> None:
    """Write upload filenames + roles so Celery workers can rehydrate session from disk."""
    session_dir = UPLOAD_ROOT / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    manifest = session_dir / "_upload_roles.json"
    payload = {"files": [{"name": f.name, "role": f.extras.get("upload_role", "auto")} for f in state.files]}
    manifest.write_text(json.dumps(payload, indent=0), encoding="utf-8")


def rehydrate_session_uploads_from_disk(session_id: str) -> bool:
    """Populate in-memory session files from UPLOAD_ROOT (for Celery workers). No-op if files already loaded."""
    state = _ensure_session(session_id)
    if state.files:
        return True
    session_dir = UPLOAD_ROOT / session_id
    if not session_dir.is_dir():
        return False
    roles_map: Dict[str, str] = {}
    manifest = session_dir / "_upload_roles.json"
    if manifest.exists():
        try:
            data = json.loads(manifest.read_text(encoding="utf-8"))
            for entry in data.get("files") or []:
                name = str(entry.get("name") or "").strip()
                if name:
                    roles_map[name] = str(entry.get("role") or "auto").strip().lower() or "auto"
        except Exception:
            roles_map = {}
    skip_exact = {
        "reviewed_contract.docx",
        "clause_analysis.csv",
        "_upload_roles.json",
        "_chat_history.json",
    }
    skip_suffixes = ("_converted.docx",)
    candidates: List[Path] = []
    for path in sorted(session_dir.iterdir()):
        if not path.is_file():
            continue
        if path.name in skip_exact or path.name.startswith("."):
            continue
        if any(path.name.endswith(suf) for suf in skip_suffixes):
            continue
        if path.suffix.lower() not in ALLOWED_EXTENSIONS:
            continue
        candidates.append(path)
    if not candidates:
        return False
    state.files = []
    for path in sorted(candidates, key=lambda p: p.stat().st_mtime):
        artifact = UploadedArtifact(name=path.name, path=path, size=path.stat().st_size)
        role = roles_map.get(path.name, "auto")
        if role not in {"auto", "contract", "supporting"}:
            role = "auto"
        artifact.extras["upload_role"] = role
        text = _extract_text(artifact.path)
        if text:
            artifact.text_preview = text[:1200]
            artifact.extras["full_text"] = text
            try:
                from .legal_preprocessor import build_legal_context

                artifact.extras["legal_context"] = build_legal_context(text)
            except Exception:
                artifact.extras["legal_context"] = {}
            artifact.extras["indexed_in_rag"] = False
        else:
            artifact.extras["image_bytes"] = _extract_images_from_pdf(artifact.path)
        state.files.append(artifact)
    try:
        _rebuild_upload_rag_index(session_id, state)
    except Exception:
        pass
    return bool(state.files)


def run_detection_analysis_for_session(session_id: str, *, message: str = "[Analyze clauses]") -> Dict[str, Any]:
    """Rebuild session from disk then run the same analysis path as the synchronous workspace."""
    if not rehydrate_session_uploads_from_disk(session_id):
        raise ValueError("No contract uploads found on disk for this session. Upload before running analysis.")
    return answer_question(session_id, message)


def reset_session(session_id: str) -> None:
    """Clear uploads and chat state."""
    _SESSION_DATA[session_id] = SessionState()
    try:
        p = _chat_history_disk_path(session_id)
        if p.is_file():
            p.unlink()
    except Exception:
        pass
    try:
        from . import config as agent_config

        if agent_config.ENABLE_VECTOR_RETRIEVER:
            from . import rag

            rag.clear_session_upload_index(session_id)
    except Exception:
        pass


def list_files(session_id: str) -> List[Dict[str, Any]]:
    state = _ensure_session(session_id)
    return [{"name": f.name, "size": f.size, "role": f.extras.get("upload_role", "auto")} for f in state.files]


def _rebuild_upload_rag_index(session_id: str, state: SessionState) -> None:
    """Rebuild upload RAG index with main/supporting doc metadata."""
    from . import config as agent_config, rag

    if not (agent_config.ENABLE_RAG and agent_config.ENABLE_VECTOR_RETRIEVER):
        return
    rag.clear_session_upload_index(session_id)
    readable = [
        (artifact, str(artifact.extras.get("full_text") or ""))
        for artifact in state.files
        if artifact.extras.get("full_text")
    ]
    if not readable:
        return
    primary_artifact = _primary_readable_artifact(state)
    primary_name = primary_artifact.name if primary_artifact is not None else max(readable, key=lambda item: len(item[1]))[0].name
    for artifact, full_text in readable:
        explicit_role = artifact.extras.get("upload_role")
        doc_type = "main" if artifact.name == primary_name or explicit_role == "contract" else "supporting"
        indexed_chunks = rag.index_uploaded_text(session_id, artifact.name, full_text, doc_type=doc_type)
        artifact.extras["indexed_in_rag"] = True
        artifact.extras["rag_doc_type"] = doc_type
        artifact.extras["rag_chunks"] = indexed_chunks


def _primary_readable_artifact(state: SessionState) -> UploadedArtifact | None:
    """Return the most likely primary contract: largest readable uploaded document."""
    contract_role = [
        artifact
        for artifact in state.files
        if artifact.extras.get("full_text") and artifact.extras.get("upload_role") == "contract"
    ]
    if contract_role:
        return max(contract_role, key=lambda artifact: len(str(artifact.extras.get("full_text") or "")))
    readable = [
        artifact
        for artifact in state.files
        if artifact.extras.get("full_text")
    ]
    if not readable:
        return None
    return max(readable, key=lambda artifact: len(str(artifact.extras.get("full_text") or "")))


def _supporting_doc_upload_warnings(state: SessionState) -> List[str]:
    """Warn at upload-time when the primary contract references missing appendices/docs."""
    readable = [
        (artifact, str(artifact.extras.get("full_text") or ""))
        for artifact in state.files
        if artifact.extras.get("full_text")
    ]
    if not readable:
        return []
    primary_artifact = _primary_readable_artifact(state)
    if primary_artifact is None:
        return []
    primary_text = str(primary_artifact.extras.get("full_text") or "")
    try:
        import agent1_clause_analyzer as agent1  # type: ignore
    except Exception:
        try:
            from . import agent1_clause_analyzer as agent1  # type: ignore
        except Exception:
            return []

    referenced = agent1.detect_referenced_supporting_docs(primary_text)
    if not referenced:
        return []

    missing = []
    for ref in referenced:
        matched = any(
            agent1.referenced_doc_matches_upload(ref.get("label", ""), artifact.name) for artifact in state.files
        )
        if not matched:
            missing.append(ref.get("label", ""))
    if not missing:
        return []
    return [
        (
            f"{primary_artifact.name}: referenced supporting documents not uploaded: "
            f"{', '.join(missing[:10])}. Upload these before analysis/export, or proceed knowing related clause outputs may be incomplete."
        )
    ]


def remove_uploaded_file(session_id: str, filename: str) -> Dict[str, Any]:
    """Remove a file from session and disk."""
    state = _ensure_session(session_id)
    keep: List[UploadedArtifact] = []
    removed = False
    for f in state.files:
        if f.name == filename:
            removed = True
            try:
                f.path.unlink(missing_ok=True)
            except Exception:
                pass
        else:
            keep.append(f)
        state.files = keep
    try:
        _rebuild_upload_rag_index(session_id, state)
    except Exception:
        pass
    try:
        _persist_upload_roles_manifest(session_id, state)
    except Exception:
        pass
    return {"removed": removed, "files": list_files(session_id)}


def _save_upload(session_id: str, upload) -> UploadedArtifact:
    """Persist an uploaded file to disk; customize to use S3 if needed."""
    suffix = Path(upload.name).suffix.lower()
    session_dir = UPLOAD_ROOT / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    target = session_dir / upload.name

    with target.open("wb") as fh:
        for chunk in upload.chunks() if hasattr(upload, "chunks") else [upload.read()]:
            fh.write(chunk)

    size = target.stat().st_size
    return UploadedArtifact(name=upload.name, path=target, size=size, mime=getattr(upload, "content_type", None))


def _extract_text(path: Path) -> str:
    """Extract text from uploaded documents with fallbacks."""
    if path.suffix.lower() == ".xlsx":
        try:
            z = zipfile.ZipFile(path)
            ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
            shared: List[str] = []
            if "xl/sharedStrings.xml" in z.namelist():
                root = ET.fromstring(z.read("xl/sharedStrings.xml"))
                for si in root.findall(f"{{{ns}}}si"):
                    shared.append("".join((t.text or "") for t in si.iter(f"{{{ns}}}t")))
            rows: List[str] = []
            sheet_names = [n for n in z.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
            for sheet in sorted(sheet_names):
                root = ET.fromstring(z.read(sheet))
                for row in root.findall(f".//{{{ns}}}row"):
                    vals: List[str] = []
                    for cell in row.findall(f"{{{ns}}}c"):
                        typ = cell.attrib.get("t")
                        value = cell.find(f"{{{ns}}}v")
                        if value is not None and value.text is not None:
                            if typ == "s" and value.text.isdigit() and int(value.text) < len(shared):
                                vals.append(shared[int(value.text)])
                            else:
                                vals.append(value.text)
                    if any(v.strip() for v in vals):
                        rows.append(" | ".join(v.strip() for v in vals if v.strip()))
            return "\n".join(rows).strip()
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("xlsx extraction failed for %s: %s", path.name, exc)
            return ""

    # Primary path: shared extractor supports PDF + DOCX + DOC with better handling.
    if shared_extract_text and path.suffix.lower() in {".pdf", ".doc", ".docx"}:
        try:
            text = shared_extract_text(str(path), include_page_markers=True)
            if text and text.strip():
                return text
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("Shared extractor failed for %s: %s", path.name, exc)

    # Secondary path: textract (optional, may be unavailable in some environments).
    if textract and path.suffix.lower() in {".pdf", ".doc", ".docx"}:
        try:
            return textract.process(str(path)).decode("utf-8", errors="ignore")
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("Textract fallback failed for %s: %s", path.name, exc)

    # DOCX fallback A: parse paragraph/cell text using python-docx.
    if path.suffix.lower() == ".docx" and DocxDocument:
        try:
            doc = DocxDocument(str(path))
            chunks: List[str] = []
            chunks.extend([p.text.strip() for p in doc.paragraphs if (p.text or "").strip()])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = (cell.text or "").strip()
                        if txt:
                            chunks.append(txt)
            joined = "\n".join(chunks).strip()
            if joined:
                return joined
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("python-docx fallback failed for %s: %s", path.name, exc)

    # DOCX fallback: extract XML text directly when shared extractor/textract fail.
    if path.suffix.lower() == ".docx" and docx2txt:
        try:
            text = docx2txt.process(str(path))
            if text and text.strip():
                return text
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("docx2txt fallback failed for %s: %s", path.name, exc)

    # PDF fallback: direct text extraction with PyMuPDF if shared extractor fails.
    if path.suffix.lower() == ".pdf" and fitz:
        try:
            parts: List[str] = []
            doc = fitz.open(str(path))
            for i, page in enumerate(doc):
                txt = page.get_text("text") or ""
                txt = txt.strip()
                if txt:
                    parts.append(f"[PAGE {i+1}]\n{txt}")
            extracted = "\n\n".join(parts).strip()
            if extracted:
                return extracted
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("PyMuPDF text fallback failed for %s: %s", path.name, exc)

    # Plaintext fallback for .txt only.
    try:
        if path.suffix.lower() in {".txt", ".csv"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        # Never decode non-txt binaries (e.g. .docx zip payload) as plain text.
        return ""
    except Exception:
        return ""


def _extract_images_from_pdf(path: Path, limit: int = 3) -> List[bytes]:
    """Fallback: render first N PDF pages to images if text extraction is weak."""
    if path.suffix.lower() != ".pdf":
        return []

    # Primary path: shared PDF renderer utility.
    if shared_render_pdf_images:
        try:
            rendered = shared_render_pdf_images(str(path), max_pages=limit, dpi=120)
            blobs = [img for _, img in rendered if img]
            if blobs:
                return blobs
        except Exception as exc:
            if DEBUG_AGENT:
                logger.warning("Shared PDF render fallback failed for %s: %s", path.name, exc)

    if not fitz:
        return []

    # Secondary path: direct PyMuPDF rendering.
    images: List[bytes] = []
    try:
        doc = fitz.open(str(path))
        for i, page in enumerate(doc):
            if i >= limit:
                break
            pix = page.get_pixmap(dpi=120)
            images.append(pix.tobytes("png"))
    except Exception:
        return []
    return images


def index_uploaded_files(session_id: str, uploads: Sequence[Any], *, upload_role: str = "auto") -> Dict[str, Any]:
    """Validate and store uploads; add your own chunking + embedding here."""
    state = _ensure_session(session_id)
    warnings: List[str] = []
    normalized_role = (upload_role or "auto").strip().lower()
    if normalized_role not in {"auto", "contract", "supporting"}:
        normalized_role = "auto"

    if len(state.files) + len(uploads) > MAX_UPLOAD_FILES:
        raise ValueError(f"Too many files. Limit: {MAX_UPLOAD_FILES}")

    for upload in uploads:
        suffix = Path(upload.name).suffix.lower()
        if DEBUG_AGENT:
            logger.info("Upload received: session=%s file=%s ext=%s", session_id, upload.name, suffix)
        if suffix not in ALLOWED_EXTENSIONS:
            warnings.append(f"Skipped {upload.name}: unsupported type.")
            if DEBUG_AGENT:
                logger.warning("Upload skipped (unsupported): file=%s ext=%s", upload.name, suffix)
            continue

        artifact = _save_upload(session_id, upload)
        artifact.extras["upload_role"] = normalized_role
        if artifact.size and artifact.size > MAX_UPLOAD_BYTES:
            warnings.append(f"Skipped {artifact.name}: exceeds {MAX_UPLOAD_BYTES // (1024*1024)} MB.")
            artifact.path.unlink(missing_ok=True)
            if DEBUG_AGENT:
                logger.warning("Upload skipped (too large): file=%s size=%s", artifact.name, artifact.size)
            continue

        text = _extract_text(artifact.path)
        if text:
            artifact.text_preview = text[:1200]
            artifact.extras["full_text"] = text
            try:
                from .legal_preprocessor import build_legal_context

                artifact.extras["legal_context"] = build_legal_context(text)
            except Exception:
                artifact.extras["legal_context"] = {}
            artifact.extras["indexed_in_rag"] = False
            if DEBUG_AGENT:
                logger.info(
                    "Upload parsed text: file=%s chars=%d preview=%s",
                    artifact.name,
                    len(text),
                    _preview(text),
                )
            try:
                from . import rag, config as agent_config

                if (
                    agent_config.ENABLE_RAG
                    and agent_config.ENABLE_VECTOR_RETRIEVER
                    and agent_config.INDEX_UPLOAD_ON_UPLOAD
                ):
                    indexed_chunks = rag.index_uploaded_text(session_id, artifact.name, text, doc_type="upload")
                    artifact.extras["indexed_in_rag"] = True
                    if DEBUG_AGENT:
                        logger.info(
                            "Upload indexed in RAG: file=%s chunks=%d session=%s",
                            artifact.name,
                            indexed_chunks,
                            session_id,
                        )
                elif agent_config.ENABLE_RAG and DEBUG_AGENT:
                    logger.info(
                        "Upload indexing deferred/bypassed: file=%s session=%s (INDEX_UPLOAD_ON_UPLOAD=0 or vector retriever disabled)",
                        artifact.name,
                        session_id,
                    )
            except Exception as exc:
                warnings.append(f"{artifact.name}: text extracted but indexing failed ({exc}).")
                if DEBUG_AGENT:
                    logger.warning("RAG indexing failed for %s: %s", artifact.name, exc)
        else:
            artifact.extras["image_bytes"] = _extract_images_from_pdf(artifact.path)
            if not artifact.extras["image_bytes"]:
                warnings.append(f"{artifact.name}: no text or images extracted; please adjust extraction logic.")
            if DEBUG_AGENT:
                logger.warning(
                    "Upload had no text: file=%s image_fallback_pages=%d",
                    artifact.name,
                    len(artifact.extras.get("image_bytes") or []),
                )

        # Same filename replaces disk in _save_upload; drop any prior in-memory entry so the UI
        # does not show duplicates (e.g. supporting multi-select or repeat submissions).
        state.files = [f for f in state.files if f.name != artifact.name]

        # TODO: build chunks + embeddings and attach to `artifact.extras` or `state.vector_index`.
        state.files.append(artifact)

    order_keys: list[str] = []
    deduped_files: dict[str, UploadedArtifact] = {}
    for f in state.files:
        if f.name not in deduped_files:
            order_keys.append(f.name)
        deduped_files[f.name] = f
    state.files = [deduped_files[k] for k in order_keys]

    try:
        _rebuild_upload_rag_index(session_id, state)
    except Exception as exc:
        warnings.append(f"Upload text extracted but RAG role indexing failed ({exc}).")
        if DEBUG_AGENT:
            logger.warning("RAG role indexing failed for session=%s: %s", session_id, exc)
    warnings.extend(_supporting_doc_upload_warnings(state))
    _persist_upload_roles_manifest(session_id, state)

    return {"files": list_files(session_id), "warnings": warnings}


DEFAULT_SYSTEM_PROMPT = (
    "You are a domain-specific assistant. Answer strictly from the provided context; "
    "if the answer is not present, say you do not have enough information. Cite file names "
    "or section labels when relevant. Keep responses concise."
)

POC_SYSTEM_PROMPT = (
    "You are an assistant for GB Aerospace supply-contract review. Use ONLY the provided context "
    "(POC knowledge document and any uploaded contract text). Do not invent policy.\n\n"
    "Scope (from POC_SCOPE): This guidance applies ONLY to supply contracts where Aerospace is supplying goods. "
    "Do not use it for purchase/procurement, sub-contracting, strategic alliances, consortiums, or other agreement types.\n\n"
    "For each of the POC clause themes (same checklist as structured analysis), structure your answer as: "
    "(1) Clause name; (2) Risk vs GB's ideal position; (3) Suggested mitigation from the document; "
    "(4) Approval for deviations (ERMC, BU Head, Legal, or Legal in consultation with Business); "
    "(5) Counterfactual: what would need to change for the clause to align with GB's position or for the finding to be acceptable.\n\n"
    "Output counterfactuals in a clearly delimited section at the end, e.g. '--- Counterfactuals ---' followed by one counterfactual per clause.\n\n"
    "Disclaimer: These evaluations are for POC validation only. Positions must be routed through Legal and must not be put into commercial practice without legal approvals."
)

POC_SCOPE_GUARD = (
    "SCOPE (POC): GB Aerospace supply contracts only — supplier supplying goods. "
    "Out of scope: procurement/purchase, sub-contracting, consortiums/strategic alliances, NDAs as the primary agreement type, "
    "and generic workflows beyond the POC playbook. Outputs are for POC validation; route through Legal before commercial use."
)

TABLE_SYSTEM_PROMPT = (
    "You are a strict contract-analysis engine for GB Aerospace. "
    + POC_SCOPE_GUARD
    + " "
    "Return ONLY valid JSON and do not include markdown fences or extra text."
)


_POC_CLAUSE_THEME_NAMES_CACHE: Optional[List[str]] = None


def poc_clause_theme_names() -> List[str]:
    """14 top-level themes aligned with Agent 1 ``CLAUSE_SPEC`` (lazy import)."""
    global _POC_CLAUSE_THEME_NAMES_CACHE
    if _POC_CLAUSE_THEME_NAMES_CACHE is None:
        from .agent1_clause_analyzer import CLAUSE_SPEC

        _POC_CLAUSE_THEME_NAMES_CACHE = [CLAUSE_SPEC[i][0] for i in sorted(CLAUSE_SPEC.keys())]
    return _POC_CLAUSE_THEME_NAMES_CACHE


_KNOWLEDGE_CACHE: Optional[dict] = None
_SUPPLY_SCOPE_SIGNALS = [
    "purchase order",
    "supply",
    "supplier",
    "goods",
    "delivery",
    "liquidated damages",
    "change order",
]
_OUT_OF_SCOPE_SIGNALS = [
    "non-disclosure agreement",
    "confidential disclosure agreement",
    "receiving party",
    "disclosing party",
    "confidentiality period",
]

CLAUSE_RULES: List[Dict[str, Any]] = [
    {
        "name": "Limitation of Liability and Exclusion of Consequential Damages",
        "keywords": ["liability", "consequential", "indirect", "cap", "%"],
        "ideal": "Liability cap at 100% of agreement/SOW value and consequential/indirect damages excluded.",
        "approval": "ERMC if liability cap exceeds 100% of contract value.",
    },
    {
        "name": "Applicable / Governing Law - Choice of Law and Jurisdiction",
        "keywords": ["governing law", "jurisdiction", "courts", "law of"],
        "ideal": "Prefer approved combinations (India/Mumbai or approved alternatives such as English law with agreed seats).",
        "approval": "BU Head when governing law/seat is outside approved options.",
    },
    {
        "name": "Dispute Resolution",
        "keywords": ["arbitration", "mediat", "siac", "lcia", "icc", "dispute"],
        "ideal": "Arbitration preferred; mediation + arbitration with defined seat/language/tribunal details.",
        "approval": "Legal Team to decide.",
    },
    {
        "name": "Firm Price",
        "keywords": ["firm price", "fixed price", "price escalation", "change in law", "raw material"],
        "ideal": "Firm/fixed pricing with controlled escalation triggers (change in law, raw material costs, delays not attributable to supplier).",
        "approval": "ERMC if term > 2 years or value > 25cr.",
    },
    {
        "name": "Force Majeure",
        "keywords": ["force majeure", "epidemic", "government order", "non-payment", "hardship"],
        "ideal": "FM includes uncontrollable events; no excuse for non-payment of delivered goods/services; economic hardship excluded.",
        "approval": "Legal team in consultation with Business Team.",
    },
    {
        "name": "Liquidated Damages",
        "keywords": ["liquidated damages", "ld", "delay", "0.5%", "5%"],
        "ideal": "LD on delayed value (not total value), with clear cap and exclusive remedy framing.",
        "approval": "ERMC based on division thresholds when LD exceeds approved percentages.",
    },
    {
        "name": "Orders Extending Beyond Termination",
        "keywords": ["termination", "expiry", "orders", "post-term", "residual period"],
        "ideal": "Termination generally ends in-effect orders unless expressly agreed; renegotiation rights for overhang orders.",
        "approval": "Legal team in consultation with Business Team.",
    },
    {
        "name": "Quantity Protection",
        "keywords": ["forecast", "quantity", "deviation", "lead time", "firm po"],
        "ideal": "Protection on quantity volatility with reimbursement beyond +/-20% and conversion to firm POs in lead time.",
        "approval": "Legal team in consultation with Business Team.",
    },
    {
        "name": "Inventory Requirements",
        "keywords": ["inventory", "raw material", "fg", "forecast", "weeks"],
        "ideal": "Inventory commitment against non-binding forecasts limited (e.g., capped weeks within lead time).",
        "approval": "Legal team in consultation with Business Team.",
    },
    {
        "name": "Change Orders Procedure",
        "keywords": ["change order", "equitable adjustment", "schedule relief", "tooling", "nre"],
        "ideal": "No changes without signed CO including price/time impact and equitable adjustment at actuals.",
        "approval": "Legal team in consultation with Business Team.",
    },
]


def _load_poc_knowledge_payload() -> dict:
    global _KNOWLEDGE_CACHE
    if _KNOWLEDGE_CACHE is not None:
        return _KNOWLEDGE_CACHE
    try:
        from . import config as agent_config
        from .knowledge_loader import load_knowledge_payload

        path = Path(agent_config.POC_KNOWLEDGE_PATH)
        if not path.exists():
            _KNOWLEDGE_CACHE = {}
            return _KNOWLEDGE_CACHE
        _KNOWLEDGE_CACHE = load_knowledge_payload(path)
    except Exception:
        _KNOWLEDGE_CACHE = {}
    return _KNOWLEDGE_CACHE


def _knowledge_as_text(payload: dict) -> str:
    """Build full knowledge text including explanation and standard positions."""
    if not payload:
        return ""
    if payload.get("raw_text"):
        return str(payload.get("raw_text"))
    from .knowledge_loader import payload_to_text
    return payload_to_text(payload)


def _extract_json_object(text: str) -> Optional[dict]:
    text = (text or "").strip()
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        return None
    try:
        return json.loads(match.group(0))
    except Exception:
        return None


def _normalize_risk(value: str) -> str:
    token = (value or "").strip().lower()
    if token in {"green", "low", "aligned"}:
        return "Green"
    if token in {"amber", "yellow", "medium", "partial"}:
        return "Amber"
    if token in {"red", "high", "misaligned", "critical"}:
        return "Red"
    return "Amber"


def _normalize_detected(value: str) -> str:
    token = (value or "").strip().lower()
    if token in {"yes", "y", "present", "found"}:
        return "Yes"
    if token in {"no", "n", "absent", "not found"}:
        return "No"
    return "Unclear"


def _looks_out_of_scope_document(contract_text: str) -> bool:
    text = (contract_text or "").lower()
    if not text:
        return False
    has_out_of_scope = any(sig in text for sig in _OUT_OF_SCOPE_SIGNALS)
    if not has_out_of_scope:
        return False
    scope_hits = 0
    for clause in CLAUSE_RULES:
        kws = clause.get("keywords") or []
        if any(kw.lower() in text for kw in kws[:3]):
            scope_hits += 1
    # NDA-like docs often contain legal terms but miss most supply-clause signals.
    return scope_hits < 3


def _sentence_candidates(text: str) -> List[str]:
    candidates = []
    for block in re.split(r"[\n\r]+", text or ""):
        block = " ".join(block.split()).strip()
        if not block:
            continue
        candidates.extend([s.strip() for s in re.split(r"(?<=[.!?])\s+", block) if s.strip()])
    return candidates


def _best_evidence(contract_text: str, keywords: List[str]) -> str:
    candidates = _sentence_candidates(contract_text)
    if not candidates:
        return ""
    best = ""
    best_score = 0
    for sent in candidates:
        lowered = sent.lower()
        score = sum(1 for kw in keywords if kw.lower() in lowered)
        if score > best_score and len(sent) >= 30:
            best_score = score
            best = sent
    return best if best_score > 0 else ""


def _out_of_scope_evidence(contract_text: str) -> str:
    for sent in _sentence_candidates(contract_text):
        lowered = sent.lower()
        if any(sig in lowered for sig in _OUT_OF_SCOPE_SIGNALS):
            return sent
    return "Document appears to be an NDA/confidentiality agreement, not a supply contract."


def _risk_for_clause(clause_name: str, evidence: str) -> tuple[str, str]:
    e = (evidence or "").lower()
    if not evidence:
        return "Amber", "Clause wording is not clearly present in the uploaded text."

    if clause_name.startswith("Limitation of Liability"):
        if "200%" in e or "unlimited" in e:
            return "Red", "Liability appears above preferred cap or uncapped."
        if "100%" in e and ("consequential" in e or "indirect" in e):
            return "Green", "Liability cap/exclusion pattern appears close to GB ideal position."
        return "Amber", "Liability text exists but cap/exclusion alignment is incomplete."

    if clause_name.startswith("Applicable / Governing Law"):
        if any(x in e for x in ["india", "mumbai", "singapore", "london", "english law"]):
            return "Green", "Governing-law/jurisdiction terms appear within commonly accepted patterns."
        return "Amber", "Governing-law terms exist but are outside preferred combinations."

    if clause_name == "Dispute Resolution":
        if "arbitration" in e:
            return "Green", "Arbitration language is present."
        if "court" in e or "litigation" in e:
            return "Amber", "Dispute wording exists but arbitration-first preference is missing."
        return "Amber", "Dispute language exists but lacks required arbitration details."

    if clause_name == "Liquidated Damages":
        if "0.5%" in e and "5%" in e:
            return "Green", "LD rate/cap pattern appears aligned."
        if "%" in e and ("ld" in e or "liquidated" in e):
            return "Amber", "LD exists but cap/rate or delayed-value basis is unclear."
        return "Amber", "LD references are incomplete."

    return "Amber", "Clause appears but needs legal validation against GB ideal wording."


def _mitigation_for_clause(clause: Dict[str, Any], risk_level: str) -> str:
    if risk_level == "Green":
        return "Retain clause wording; confirm final legal drafting and cross-clause consistency."
    if risk_level == "Red":
        return f"Replace with GB ideal language: {clause['ideal']}"
    return (
        f"Partially align clause to GB ideal language: {clause['ideal']} "
        "and route deviation through the stated approval path."
    )


def _brief_counterfactual_narrative(row: Dict[str, Any]) -> str:
    """Produce a SHORT 'what-if' statement for counterfactual display. Legal tone, no full GB ideal text.
    Clause-specific first to avoid wrong what-if (e.g. LD text for Dispute/FM).
    """
    risk_level = (row.get("risk_level") or "Amber").strip()
    shift = (row.get("expected_risk_shift") or ("Amber -> Green" if risk_level != "Green" else "Maintain Green")).strip()
    rationale = (row.get("risk_rationale") or row.get("risk_trigger") or "").strip()
    clause_name = (row.get("clause_name") or "").lower()
    r_lower = rationale.lower()

    if risk_level == "Green":
        return f"If the final negotiated text preserves the current structure, risk should remain Green."

    # Clause-specific what-if first (robust: never apply LD what-if to Dispute/FM, etc.)
    if "dispute resolution" in clause_name:
        return f"If arbitration-first structure with seat/language/rules is adopted, risk can shift {shift}."
    if "force majeure" in clause_name:
        return f"If FM clause excludes economic hardship and preserves payment for delivered goods, risk can shift {shift}."
    if "liquidated damages" in clause_name or "ld" in clause_name:
        return f"If LD is recast on delayed value with ~0.5%/week and ~5% cap, risk can shift {shift}."
    if "limitation of liability" in clause_name or "lod" in clause_name:
        return f"If liability cap is set to 100% and consequential damages excluded, risk can shift {shift}."
    if "governing law" in clause_name or "applicable" in clause_name:
        return f"If governing law is updated to an approved combination (e.g. India/Mumbai), risk can shift {shift}."
    if "firm price" in clause_name:
        return f"If escalation is limited to change order, change in law, and raw material costs, risk can shift {shift}."
    if "orders extending" in clause_name:
        return f"If post-termination orders require repricing or mutual addendum, risk can shift {shift}."
    if "quantity protection" in clause_name:
        return f"If forecast deviations beyond +/-20% trigger reimbursement at actuals, risk can shift {shift}."
    if "inventory" in clause_name:
        return f"If inventory cap is set to max 4 weeks once forecast enters lead time, risk can shift {shift}."
    if "change order" in clause_name:
        return f"If changes require a signed CO with price/time impact, risk can shift {shift}."
    if "aerospace" in clause_name:
        return f"If aerospace terms (scope, specs, warranties) align to GB playbook, risk can shift {shift}."

    # Fallback: rationale-based patterns (only when clause not matched above)
    if "approved combination" in r_lower or "governing law" in r_lower:
        return f"If governing law is updated to an approved combination (e.g. India/Mumbai), risk can shift {shift}."
    if "payment" in r_lower and ("carve-out" in r_lower or "economic" in r_lower):
        return f"If FM clause excludes economic hardship and preserves payment obligations for delivered goods, risk can shift {shift}."
    if "termination" in r_lower and ("repricing" in r_lower or "pricing" in r_lower):
        return f"If post-termination orders require repricing or mutual addendum, risk can shift {shift}."
    if "forecast" in r_lower or "reimbursement" in r_lower or "lead-time" in r_lower:
        return f"If forecast deviations beyond +/-20% trigger reimbursement at actuals, risk can shift {shift}."
    if "change order" in r_lower or "signature" in r_lower or "equitable adjustment" in r_lower:
        return f"If changes require a signed CO with price/time impact, risk can shift {shift}."

    # Derive from rationale
    MAX_SUMMARY = 70
    if rationale:
        first = rationale.split(".")[0].split(";")[0].strip()
        skip_phrases = ["add ", "deviation", "approval", "clause text", "clause present", "clause needs"]
        if not any(p in first.lower() for p in skip_phrases) and len(first) > 15:
            if len(first) > MAX_SUMMARY:
                first = first[:MAX_SUMMARY].rsplit(" ", 1)[0] + "..."
            return f"If {first}, risk can shift {shift}."
    return f"If the clause is aligned to GB policy, risk can shift {shift}."


def _counterfactual_for_clause(clause: Dict[str, Any], risk_level: str, evidence: str, out_of_scope: bool) -> str:
    if out_of_scope:
        return (
            "If this were converted into a supply-contract format with explicit wording for this clause, "
            "the tool could perform a clause-level compliant assessment."
        )
    if risk_level == "Green":
        return "If final negotiated text preserves this structure without adverse carve-outs, this finding should remain acceptable."
    if risk_level == "Red":
        return f"If the clause is replaced with GB ideal language ({clause['ideal']}), risk can reduce from Red to Amber/Green."
    if evidence:
        return "If missing parameters (cap/seat/cap percentage/approval trigger) are explicitly added to match GB ideal wording, risk can reduce to Green."
    return "If explicit clause text is inserted in line with GB ideal wording, this row can move from Unclear/Amber to Green."


def _build_clause_rows(contract_text: str, out_of_scope: bool) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    out_scope_snippet = _out_of_scope_evidence(contract_text) if out_of_scope else ""
    for idx, clause in enumerate(CLAUSE_RULES, start=1):
        evidence = _best_evidence(contract_text, clause["keywords"])
        detected = "Yes" if evidence else "Unclear"
        risk_level, risk_rationale = _risk_for_clause(clause["name"], evidence)

        if out_of_scope:
            detected = "Unclear"
            risk_level = "Amber"
            evidence = out_scope_snippet
            risk_rationale = (
                "Document type appears out-of-scope for this POC (supply-contract clauses). "
                "Route for legal review before relying on this assessment."
            )

        rows.append(
            {
                "clause_name": clause["name"],
                "detected": detected,
                "uploaded_position": (
                    "Out-of-scope document (NDA/confidentiality format) for this supply-contract POC."
                    if out_of_scope
                    else (evidence or "Insufficient evidence in uploaded text.")
                ),
                "gb_ideal_position": clause["ideal"],
                "risk_level": risk_level,
                "risk_rationale": risk_rationale,
                "mitigation_recommendation": _mitigation_for_clause(clause, risk_level),
                "approval_path": clause["approval"],
                "evidence_snippet": evidence or "No direct matching clause text found.",
                "evidence_quote": (evidence or "").strip(),
                "evidence_source": "primary",
                "evidence_location_hint": "",
                "anchoring_warning": "",
                "knowledge_reference": (
                    "POC scope limitation - supply contracts only"
                    if out_of_scope
                    else f"POC clause {idx}: {clause['name']}"
                ),
                "counterfactual": _counterfactual_for_clause(clause, risk_level, evidence, out_of_scope),
            }
        )
    return rows


def _to_agent1_style_markdown(clause_table: List[Dict[str, str]]) -> str:
    lines = [
        "| Clause | Requirement | Met/Gap | References |",
        "|--------|-------------|---------|------------|",
    ]
    for idx, row in enumerate(clause_table, start=1):
        clause = row.get("clause_name", "")
        requirement = row.get("gb_ideal_position", "")
        risk = row.get("risk_level", "Amber")
        status = row.get("detected", "Unclear")
        rationale = row.get("risk_rationale", "")
        references = row.get("knowledge_reference", "")
        met_gap = f"{status} ({risk}) - {rationale}"
        lines.append(f"| {idx}. {clause} | {requirement} | {met_gap} | {references} |")
    return "\n".join(lines)


def _agent3_to_counterfactuals(agent3_markdown: str) -> str:
    def _compact(text: str, limit: int = 220) -> str:
        cleaned = " ".join((text or "").split()).strip()
        if not cleaned:
            return ""
        first = re.split(r"(?<=[.!?])\s+", cleaned)[0].strip()
        base = first or cleaned
        if len(base) <= limit:
            return base
        return base[:limit].rstrip() + "..."

    lines = [ln.strip() for ln in (agent3_markdown or "").splitlines() if ln.strip().startswith("|")]
    bullets: List[str] = []
    for ln in lines[2:]:  # skip table header + separator
        parts = [p.strip() for p in ln.strip("|").split("|")]
        if len(parts) < 3:
            continue
        clause, _risk, mitigation = parts[:3]
        if clause.lower().startswith("no risks identified"):
            continue
        if mitigation in {"-", "—", ""}:
            continue
        compact = _compact(mitigation)
        if compact:
            bullets.append(f"- {clause}: {compact}")
    if len(bullets) < 2:
        return ""
    return "\n".join(bullets).strip()


def _fallback_mitigation_checklist_from_clause_table(clause_table: List[Dict[str, str]]) -> str:
    lines = [
        "🔍 **Risk Mitigation Checklist**",
        "",
        "| Clause | Risk | Mitigation Recommendation |",
        "|--------|------|---------------------------|",
    ]
    risky_rows = []
    for idx, row in enumerate(clause_table or [], start=1):
        risk = (row.get("risk_level") or "Amber").strip()
        detected = (row.get("detected") or "Unclear").strip()
        if risk == "Green" and detected == "Yes":
            continue
        risky_rows.append(
            (
                f"{idx}. {row.get('clause_name', '')}".strip(),
                row.get("risk_rationale", "") or "Risk requires legal review.",
                row.get("mitigation_recommendation", "") or "Route to Legal for clause-specific fallback language.",
            )
        )
    if not risky_rows:
        lines.append("| No risks identified | - | - |")
        return "\n".join(lines)
    for clause, risk_text, mitigation in risky_rows:
        lines.append(f"| {clause} | {risk_text} | {mitigation} |")
    return "\n".join(lines)


def _parse_checklist_rows(markdown_text: str) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for raw_line in (markdown_text or "").splitlines():
        line = raw_line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|\s*-+\s*\|\s*-+\s*\|\s*-+\s*\|?$", line):
            continue
        parts = [p.strip() for p in line.strip("|").split("|")]
        if len(parts) < 3:
            continue
        clause, risk, mitigation = parts[:3]
        if clause.lower() == "clause":
            continue
        if clause.lower().startswith("no risks identified"):
            continue
        rows.append(
            {
                "clause": clause,
                "risk": risk or "Risk requires legal review.",
                "mitigation": mitigation or "Route to Legal for clause-specific mitigation.",
            }
        )
    return rows


def _checklist_rows_from_clause_table(clause_table: List[Dict[str, str]]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for idx, row in enumerate(clause_table or [], start=1):
        risk_level = (row.get("risk_level") or "Amber").strip()
        detected = (row.get("detected") or "Unclear").strip()
        if risk_level == "Green" and detected == "Yes":
            continue
        priority = "Medium"
        if risk_level == "Red":
            priority = "High"
        elif risk_level == "Amber":
            priority = "Medium"
        elif risk_level == "Green":
            priority = "Low"
        rows.append(
            {
                "clause": f"{idx}. {row.get('clause_name', '').strip()}",
                "risk_trigger": row.get("risk_trigger", "") or row.get("risk_rationale", "") or "Risk requires legal review.",
                "risk": row.get("risk_rationale", "") or "Risk requires legal review.",
                "mitigation": row.get("mitigation_recommendation", "") or "Route to Legal for clause-specific mitigation.",
                "approval_path": row.get("approval_path", "") or "Legal team in consultation with Business Team.",
                "priority": priority,
            }
        )
    return rows


def _counterfactual_rows_from_clause_table(clause_table: List[Dict[str, str]]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for idx, row in enumerate(clause_table or [], start=1):
        clause_name = row.get("clause_name", "").strip() or f"Clause {idx}"
        risk_level = (row.get("risk_level") or "Amber").strip()
        evidence = (row.get("evidence_snippet") or "").strip()
        trigger = (row.get("risk_trigger") or row.get("risk_rationale") or "").strip()
        shift = (row.get("expected_risk_shift") or ("Amber -> Green" if risk_level != "Green" else "Maintain Green")).strip()
        current = evidence.split(" | ")[0].strip() if evidence else "No direct clause evidence captured."
        if len(current) > 120:
            current = current[:120] + "..."
        reason = trigger or "Clause needs legal alignment against policy controls."
        # Brief "what-if" narrative only; Agent 3 checklist keeps full mitigation text
        what_if = _brief_counterfactual_narrative(row)
        rows.append(
            {
                "clause": f"{idx}. {clause_name}",
                "current_text": current,
                "required_change": what_if,
                "risk_shift": shift,
                "reason": reason,
            }
        )
    return rows


def _render_counterfactuals(counterfactual_rows: List[Dict[str, str]]) -> str:
    lines: List[str] = []
    for row in counterfactual_rows:
        what_if = row.get("required_change", "")
        if what_if:
            lines.append(f"- {row.get('clause', '')}: {what_if}")
    return "\n".join(lines).strip()


def _finalize_clause_table_citations(
    clause_table: List[Dict[str, Any]],
    *,
    primary_text: str,
    supporting_doc_texts: Dict[str, str],
) -> Dict[str, Any]:
    """Ensure citation columns exist; verify quotes for rows not already checked in Agent1."""
    trace = {"evidence_quote_verification": {"passed": 0, "failed": 0}}
    if not clause_table:
        return trace
    try:
        from .contract_search import verify_evidence_quote
    except Exception:
        return trace

    for row in clause_table:
        if not row.get("evidence_quote"):
            ev = (row.get("evidence_snippet") or "").strip()
            first = ev.split(" | ")[0].split("[Source:")[0].split("[Definitions:")[0].strip()
            if first and not re.match(r"^no direct matching", first, re.IGNORECASE):
                row["evidence_quote"] = first[:1200]
        if not row.get("evidence_source"):
            row["evidence_source"] = "primary"
        if row.get("_quote_verified_by_agent1") or row.get("anchoring_warning"):
            continue
        q = (row.get("evidence_quote") or "").strip()
        det = (row.get("detected") or "").strip()
        if det != "Yes" or not q:
            continue
        ok, src = verify_evidence_quote(q, primary_text=primary_text or "", supporting_texts=supporting_doc_texts)
        if ok:
            trace["evidence_quote_verification"]["passed"] += 1
            if src and src != "primary":
                row["evidence_source"] = src
        else:
            trace["evidence_quote_verification"]["failed"] += 1
            row["detected"] = "Unclear"
            row["anchoring_warning"] = row.get("anchoring_warning") or (
                "Evidence quote could not be verified as a verbatim substring of uploaded contract/supporting text."
            )
    return trace


def _sources_used_strip(clause_table: List[Dict[str, Any]], *, max_items: int = 4) -> List[Dict[str, str]]:
    """Compact excerpts for optional UI (no library jargon)."""
    out: List[Dict[str, str]] = []
    for row in clause_table or []:
        q = (row.get("evidence_quote") or "").strip()
        if not q:
            continue
        name = (row.get("clause_name") or "")[:120]
        src = (row.get("evidence_source") or "primary")[:120]
        clip = " ".join(q.split())
        if len(clip) > 240:
            clip = clip[:237] + "…"
        out.append({"clause": name, "source": src, "excerpt": clip})
        if len(out) >= max_items:
            break
    return out


def _search_windows_for_llm(contract_text: str, theme_names: List[str]) -> str:
    """Keyword search windows before LLM synthesis (Mike-style find-then-assert)."""
    try:
        from .contract_search import collect_keyword_spans, extract_windows, format_search_windows_for_prompt
    except Exception:
        return ""
    blocks: List[str] = []
    for name in theme_names:
        kws: List[str] = []
        for rule in CLAUSE_RULES:
            if rule.get("name") == name:
                kws = list(rule.get("keywords") or [])
                break
        if not kws and name:
            kws = [name.split()[0]] if name.split() else []
        spans = collect_keyword_spans(contract_text or "", kws[:8])
        wins = extract_windows(contract_text or "", spans, radius=380, max_windows=2)
        if wins:
            blocks.append(f"[{name}]\n" + "\n".join(wins))
    blob = "\n\n".join(blocks).strip()
    if not blob:
        return ""
    return format_search_windows_for_prompt([blob[:5500]])


def _extract_clause_table_via_llm(
    *,
    model_id: str,
    contract_text: str,
    user_question: str,
    knowledge_text: str,
) -> List[Dict[str, str]]:
    from . import config as agent_config
    from .local_llm import call_local_chat

    names = poc_clause_theme_names()
    clause_list = "\n".join([f"- {name}" for name in names])
    n_clauses = len(names)
    search_block = _search_windows_for_llm(contract_text or "", names)
    prompt = f"""
You must evaluate the uploaded contract against exactly these {n_clauses} clause themes (in order):
{clause_list}

{POC_SCOPE_GUARD}

User question:
{user_question}

POC knowledge:
{_sample_text_for_context(knowledge_text, 12000, label="POC Knowledge")}

Deterministic search windows (keyword hits; use to locate language; still cite verbatim from contract):
{search_block or "[No keyword windows extracted — rely on contract sample below.]"}

Uploaded contract text:
{_sample_text_for_context(contract_text, 18000, label="Contract")}

Return STRICT JSON object with this schema:
{{
  "rows": [
    {{
      "clause_name": "...",
      "detected": "Yes|No|Unclear",
      "uploaded_position": "...",
      "gb_ideal_position": "...",
      "risk_level": "Green|Amber|Red",
      "risk_rationale": "...",
      "mitigation_recommendation": "...",
      "approval_path": "...",
      "evidence_snippet": "...",
      "evidence_quote": "<verbatim substring copy-pasteable from uploaded contract or blank>",
      "evidence_source": "primary or supporting filename",
      "evidence_location_hint": "optional section or page reference",
      "knowledge_reference": "...",
      "counterfactual": "..."
    }}
  ]
}}

Rules:
1) Include exactly {n_clauses} rows (one per clause from the list, in order).
2) If missing evidence, use detected="No" or "Unclear" and explain.
3) evidence_quote MUST be empty unless it is an exact contiguous substring of the uploaded contract text (after normalizing only whitespace). No paraphrase. If unsure, leave evidence_quote blank and lower confidence in rationale.
4) evidence_snippet may restate or shorten but should align with evidence_quote when both are present.
5) gb_ideal_position and mitigation must track the POC knowledge document wording; do not substitute generic legal boilerplate.
6) For redlines: keep suggested minimal edits conceptually — change the smallest span needed when you describe adjustments in uploaded_position or counterfactual (do not rewrite whole articles unless necessary).
7) Do not invent policy beyond the provided POC knowledge.
8) Output JSON only.
""".strip()

    raw = call_local_chat(
        prompt,
        system_message=TABLE_SYSTEM_PROMPT,
        model_id=model_id,
        temperature=getattr(agent_config, "LOCAL_LLM_TEMPERATURE_EXTRACTION", 0.0),
        image_bytes=None,
    )
    payload = _extract_json_object(raw) or {}
    rows = payload.get("rows", []) if isinstance(payload, dict) else []
    if not isinstance(rows, list):
        return []

    normalized: List[Dict[str, str]] = []
    for idx, clause_name in enumerate(poc_clause_theme_names()):
        row = rows[idx] if idx < len(rows) and isinstance(rows[idx], dict) else {}
        ev_snip = str(row.get("evidence_snippet") or "No direct matching clause text found.")
        ev_quote = str(row.get("evidence_quote") or "").strip()
        if not ev_quote:
            ev_quote = ev_snip.split(" | ")[0].strip()[:1200]
        normalized.append(
            {
                # Force fixed clause schema for consistent table semantics.
                "clause_name": clause_name,
                "detected": _normalize_detected(str(row.get("detected") or "Unclear")),
                "uploaded_position": str(row.get("uploaded_position") or "Insufficient evidence in uploaded text."),
                "gb_ideal_position": str(row.get("gb_ideal_position") or "See POC knowledge reference."),
                "risk_level": _normalize_risk(str(row.get("risk_level") or "Amber")),
                "risk_rationale": str(row.get("risk_rationale") or "Limited evidence to assess risk confidently."),
                "mitigation_recommendation": str(row.get("mitigation_recommendation") or "Route to Legal for clause-specific fallback language."),
                "approval_path": str(row.get("approval_path") or "Legal team in consultation with Business Team"),
                "evidence_snippet": ev_snip,
                "evidence_quote": ev_quote,
                "evidence_source": str(row.get("evidence_source") or "primary"),
                "evidence_location_hint": str(row.get("evidence_location_hint") or ""),
                "anchoring_warning": "",
                "knowledge_reference": str(row.get("knowledge_reference") or f"POC clause {idx + 1}"),
                "counterfactual": str(row.get("counterfactual") or "Add explicit clause language aligned to GB ideal position."),
            }
        )
    return normalized


def answer_question(
    session_id: str,
    message: str,
    *,
    top_k: int = 3,
) -> Dict[str, Any]:
    """Route the user question through RAG (POC knowledge + uploads) and local LLM."""
    state = _ensure_session(session_id)
    _load_chat_history_from_disk_if_empty(session_id, state)
    message = (message or "").strip()
    if not message:
        raise ValueError("Message cannot be empty.")

    from . import config as agent_config

    if agent_config.DEBUG_AGENT:
        logger.info(
            "Ask received: session=%s message_chars=%d uploads_in_session=%d rag_enabled=%s",
            session_id,
            len(message),
            len(state.files),
            agent_config.ENABLE_RAG,
        )
        logger.info(
            "Provider selected: provider=%s bedrock_enabled=%s local_model=%s bedrock_model=%s",
            agent_config.LLM_PROVIDER,
            agent_config.ENABLE_BEDROCK,
            agent_config.LOCAL_LLM_MODEL,
            agent_config.BEDROCK_MODEL_ID,
        )
        logger.debug("Ask message preview: %s", _preview(message))

    context_snippets: List[str] = []
    for f in state.files:
        if f.extras.get("full_text"):
            full = str(f.extras["full_text"])
            sampled = _sample_text_for_context(full, 6000, label=f.name)
            context_snippets.append(f"File: {f.name}\n{sampled}")
        elif f.text_preview:
            context_snippets.append(f"File: {f.name}\n{f.text_preview}")

    knowledge_payload = _load_poc_knowledge_payload()
    knowledge_text = _knowledge_as_text(knowledge_payload)
    if knowledge_text:
        sampled_knowledge = _sample_text_for_context(knowledge_text, 8000, label="POC Knowledge")
        context_snippets.append(f"POC Knowledge Document:\n{sampled_knowledge}")

    # If no text context, fall back to the first available image bytes (local LLM may ignore).
    image_payload: List[bytes] = []
    for f in state.files:
        image_payload = f.extras.get("image_bytes") or []
        if image_payload:
            break

    prompt_parts = [
        "Context (POC positions and/or uploaded documents):",
        "\n---\n".join(context_snippets) if context_snippets else "[No context available. Upload documents or ensure POC knowledge file exists.]",
        "\n\nUser question:",
        message,
    ]
    raw_prompt = "\n".join(prompt_parts)
    if agent_config.DEBUG_AGENT:
        logger.info(
            "Prompt built: context_snippets=%d prompt_chars=%d image_payloads=%d",
            len(context_snippets),
            len(raw_prompt),
            len(image_payload),
        )
        logger.debug("Prompt preview: %s", _preview(raw_prompt))

    # Keep this request text for logs/diagnostics; final output is evidence-first deterministic.
    reply = ""
    main_answer = ""
    counterfactuals_text = ""
    clause_table: List[Dict[str, str]] = []
    agent2_output = ""
    agent3_output = ""
    orchestration_meta: Dict[str, Any] = {}
    checklist_rows: List[Dict[str, str]] = []
    counterfactual_rows: List[Dict[str, str]] = []
    executive_read_payload: Optional[Dict[str, Any]] = None
    paragraph_index: Optional[List[Dict[str, Any]]] = None

    # Build structured table output deterministically from extracted evidence.
    analysis_warnings: List[str] = []

    # Separate uploaded files into: (a) primary contract text (largest / first readable doc)
    # and (b) supporting docs. All file texts enrich the analysis but are passed distinctly
    # so the agent can correctly attribute references to the right document.
    all_file_texts: List[tuple] = [
        (f.name, str(f.extras.get("full_text") or ""))
        for f in state.files
        if f.extras.get("full_text")
    ]
    uploaded_filenames_all = [f.name for f in state.files]

    # Primary contract: explicit contract upload first, then largest readable file.
    primary_artifact = _primary_readable_artifact(state)
    if primary_artifact is not None:
        primary_name = primary_artifact.name
        primary_text = str(primary_artifact.extras.get("full_text") or "")
        supporting_doc_texts: Dict[str, str] = {
            f.name: str(f.extras.get("full_text") or "")
            for f in state.files
            if f.name != primary_name and f.extras.get("full_text")
        }
    else:
        primary_name, primary_text = "", ""
        supporting_doc_texts = {}

    # Full combined text (for backward-compat functions that need one blob)
    uploaded_full_text = "\n\n".join(txt for _, txt in all_file_texts)

    if uploaded_full_text and knowledge_text:
        try:
            out_of_scope = _looks_out_of_scope_document(uploaded_full_text)
            clause_table = []
            agent2_output = ""
            agent3_output = ""
            try:
                from .master_orchestrator import run_contract_analysis_master

                clause_table, agent2_output, agent3_output, orch_warns, orchestration_meta = (
                    run_contract_analysis_master(
                        uploaded_full_text=uploaded_full_text,
                        primary_text=primary_text or uploaded_full_text,
                        primary_name=primary_name,
                        supporting_doc_texts=supporting_doc_texts,
                        uploaded_filenames_all=uploaded_filenames_all,
                        knowledge_payload=knowledge_payload,
                        rag_session_id=session_id,
                    )
                )
                analysis_warnings.extend(orch_warns)
            except Exception as exc:
                if agent_config.DEBUG_AGENT:
                    logger.warning("Master orchestrator failed before table fallbacks: %s", exc)
            if not clause_table:
                clause_table = _build_clause_rows(uploaded_full_text, out_of_scope=out_of_scope)
                agent2_output = ""
                agent3_output = ""
            if clause_table:
                citation_trace = _finalize_clause_table_citations(
                    clause_table,
                    primary_text=primary_text or uploaded_full_text,
                    supporting_doc_texts=supporting_doc_texts,
                )
                _trace_row0 = clause_table[0] if clause_table else {}
                _agent_trace = dict(_trace_row0.pop("_analysis_trace", None) or {})
                _v0 = dict(_agent_trace.pop("evidence_quote_verification", None) or {})
                _v1 = citation_trace.get("evidence_quote_verification") or {}
                _agent_trace["evidence_quote_verification"] = {
                    "passed": int(_v0.get("passed", 0)) + int(_v1.get("passed", 0)),
                    "failed": int(_v0.get("failed", 0)) + int(_v1.get("failed", 0)),
                }
                _agent_trace["sources_used"] = _sources_used_strip(clause_table)
                orchestration_meta["analysis_trace"] = _agent_trace

                if out_of_scope:
                    analysis_warnings.append(
                        "Uploaded document appears to be out-of-scope (e.g., NDA/confidentiality format). "
                        "Results are best-effort and should not be treated as supply-contract clause conclusions."
                    )

                # Surface missing supporting documents as analysis warnings
                first_row = clause_table[0] if clause_table else {}
                missing_docs = first_row.pop("_missing_supporting_docs", None) or []
                referenced_docs = first_row.pop("_referenced_supporting_docs", None) or []
                if referenced_docs:
                    missing_labels = [d["label"] for d in missing_docs]
                    if missing_labels:
                        doc_list = ", ".join(missing_labels[:8])
                        analysis_warnings.append(
                            f"Supporting documents referenced in the contract but NOT uploaded: {doc_list}. "
                            "Upload these documents for a more accurate and complete analysis. "
                            "Missing documents may affect evidence quality for specific clauses."
                        )
                    uploaded_labels = [d["label"] for d in referenced_docs if d.get("uploaded")]
                    if uploaded_labels:
                        analysis_warnings.append(
                            f"Supporting documents detected and available for analysis: {', '.join(uploaded_labels[:5])}."
                        )

                red_count = sum(1 for row in clause_table if row.get("risk_level") == "Red")
                amber_count = sum(1 for row in clause_table if row.get("risk_level") == "Amber")
                green_count = sum(1 for row in clause_table if row.get("risk_level") == "Green")
                n_positions = len(clause_table)
                n_review = sum(
                    1
                    for row in clause_table
                    if not (
                        (row.get("risk_level") or "").strip() == "Green"
                        and (row.get("detected") or "").strip() == "Yes"
                    )
                )
                main_answer = (
                    f"Clause analysis completed across **{n_positions}** POC positions "
                    f"({n_review} requiring attention or clarification, "
                    f"{n_positions - n_review} aligned with baseline where applicable). "
                    f"Risk mix: Red={red_count}, Amber={amber_count}, Green={green_count}. "
                    "Use **Risk Analysis and Details**, the mitigation checklist, and counterfactuals below; "
                    "export CSV/DOCX for Legal."
                )
                counterfactual_rows = _counterfactual_rows_from_clause_table(clause_table)
                counterfactuals_text = _render_counterfactuals(counterfactual_rows)

                # Agent 2/3 from orchestrator; fallback if orchestrator didn't return them
                if not agent2_output or not agent3_output:
                    try:
                        import agent2_reviewer as agent2  # type: ignore
                        import agent3_mitigation_checklist as agent3  # type: ignore

                        agent1_output = _to_agent1_style_markdown(clause_table)
                        if not agent2_output:
                            agent2_output = agent2.generate_risk_mitigation(
                                agent1_output=agent1_output,
                                po_text=primary_text or uploaded_full_text,
                                terms_text=knowledge_text,
                                clause_table=clause_table,
                            )
                        if not agent3_output:
                            if hasattr(agent3, "generate_mitigation_checklist_from_table"):
                                agent3_output = agent3.generate_mitigation_checklist_from_table(clause_table)
                            else:
                                agent3_output = agent3.generate_mitigation_checklist(agent2_output)
                    except Exception as exc:
                        if agent_config.DEBUG_AGENT:
                            logger.warning("Agent2/3 fallback skipped: %s", exc)
                if "No risks identified" in (agent3_output or ""):
                    has_non_green = any((row.get("risk_level") or "Amber") != "Green" for row in clause_table)
                    has_not_yes = any((row.get("detected") or "Unclear") != "Yes" for row in clause_table)
                    if has_non_green or has_not_yes:
                        agent3_output = _fallback_mitigation_checklist_from_clause_table(clause_table)
                checklist_rows = _checklist_rows_from_clause_table(clause_table)
                if clause_table and not (agent2_output or "").strip():
                    analysis_warnings.append(
                        "Risk narrative (Agent 2) produced no text. Check Ollama/Bedrock and LOCAL_LLM_MODEL_CHAT; "
                        "clause table, checklist, and exports below still reflect Agent 1."
                    )
                if getattr(agent_config, "ENABLE_EXECUTIVE_READ", True) and clause_table:
                    try:
                        from .executive_read import generate_executive_read_bundle

                        n_exec = int(getattr(agent_config, "EXECUTIVE_READ_MAX_CLAUSES", 10))
                        rows_slice = list(clause_table[:n_exec])
                        supp_bits: List[str] = []
                        for name, txt in list(supporting_doc_texts.items())[:6]:
                            frag = " ".join(str(txt or "").split())[:420]
                            if frag:
                                supp_bits.append(f"{name}: {frag}")
                        supp_summary = "\n".join(supp_bits)[:4000]
                        text_for_index = (primary_text or uploaded_full_text or "").strip()
                        bundle = generate_executive_read_bundle(
                            clause_table=rows_slice,
                            primary_text=text_for_index,
                            supporting_summary=supp_summary,
                            knowledge_excerpt=_sample_text_for_context(knowledge_text, 5000, label="POC Knowledge"),
                        )
                        executive_read_payload = bundle.get("executive_read")
                        paragraph_index = bundle.get("paragraph_index")
                    except Exception as exc:
                        if agent_config.DEBUG_AGENT:
                            logger.warning("Executive read skipped: %s", exc)
                        executive_read_payload = None
                        paragraph_index = None
            if not checklist_rows and clause_table:
                checklist_rows = _checklist_rows_from_clause_table(clause_table)
                if checklist_rows and not agent3_output:
                    agent3_output = _fallback_mitigation_checklist_from_clause_table(clause_table)
            if not counterfactual_rows and clause_table:
                counterfactual_rows = _counterfactual_rows_from_clause_table(clause_table)
                counterfactuals_text = _render_counterfactuals(counterfactual_rows)
        except Exception as exc:
            if agent_config.DEBUG_AGENT:
                logger.warning("Structured table extraction failed: %s", exc)
    elif not uploaded_full_text:
        main_answer = "No uploaded contract text available to analyze. Please upload a readable document."
    else:
        main_answer = "POC knowledge source is unavailable. Please verify POC knowledge file configuration."

    if agent_config.DEBUG_AGENT:
        logger.info(
            "Parse complete: reply_chars=%d has_counterfactuals_marker=%s main_chars=%d counterfactual_chars=%d",
            len(reply or ""),
            "--- Counterfactuals ---" in reply,
            len(main_answer or ""),
            len(counterfactuals_text or ""),
        )
        logger.debug("Main answer preview: %s", _preview(main_answer))
        logger.debug("Counterfactual preview: %s", _preview(counterfactuals_text))

    # Track chat history (include counterfactuals in assistant entry for separate display).
    state.history.append({"role": "user", "content": message})
    state.history.append({
        "role": "assistant",
        "content": main_answer,
        "counterfactuals": counterfactuals_text or None,
        "clause_table": clause_table,
        "risk_details_markdown": agent2_output or None,
        "mitigation_checklist_markdown": agent3_output or None,
        "mitigation_checklist_rows": checklist_rows,
        "counterfactual_rows": counterfactual_rows,
        "executive_read": executive_read_payload,
        "paragraph_index": paragraph_index,
        "copy_ready": True,
        "orchestration_meta": orchestration_meta if orchestration_meta else None,
        "analysis_trace": orchestration_meta.get("analysis_trace") if orchestration_meta else None,
    })
    if len(state.history) > HISTORY_LIMIT:
        state.history = state.history[-HISTORY_LIMIT:]
    _persist_chat_history_to_disk(session_id, state)

    return {
        "answer": main_answer,
        "counterfactuals": counterfactuals_text or None,
        "warnings": analysis_warnings,
        "risk_details_markdown": agent2_output or None,
        "mitigation_checklist_markdown": agent3_output or None,
        "mitigation_checklist_rows": checklist_rows,
        "counterfactual_rows": counterfactual_rows,
        "copy_ready": True,
        "chat_history": list(state.history),
        "used_files": [f.name for f in state.files],
        "executive_read": executive_read_payload,
        "paragraph_index": paragraph_index,
        "orchestration_meta": orchestration_meta if orchestration_meta else None,
        "analysis_trace": orchestration_meta.get("analysis_trace") if orchestration_meta else None,
    }


def _filter_edit_instructions_for_hitl(
    all_instructions: List[Dict[str, str]],
    *,
    accepted_clause_ids: Optional[Sequence[str]],
    require_hitl: bool,
) -> List[Dict[str, str]]:
    """When require_hitl is True, keep only rows whose clause_id is in accepted_clause_ids."""
    if not require_hitl:
        if accepted_clause_ids is None:
            return list(all_instructions)
        acc = {str(x) for x in accepted_clause_ids}
        return [i for i in all_instructions if i.get("clause_id") in acc]
    acc = {str(x) for x in (accepted_clause_ids or [])}
    return [i for i in all_instructions if i.get("clause_id") in acc]


def _validate_redline_template_source(
    primary_artifact: Optional[UploadedArtifact],
    source_docx_path: Optional[Path],
    session_id: str,
) -> None:
    """Redlines must target only the primary contract DOCX, never supporting uploads."""
    if primary_artifact is None or source_docx_path is None:
        return
    prim = primary_artifact.path.resolve()
    src = Path(source_docx_path).resolve()
    suf = prim.suffix.lower()
    if suf == ".docx":
        if src != prim:
            raise ValueError(
                "Redlined DOCX must be generated only from the primary contract file; "
                "supporting documents are read-only context."
            )
    elif suf == ".pdf":
        expected = (UPLOAD_ROOT / session_id / f"{prim.stem}_converted.docx").resolve()
        if src != expected:
            raise ValueError(
                "Redlined DOCX must be converted only from the primary contract PDF."
            )


def generate_reviewed_contract_docx(
    session_id: str,
    *,
    accepted_clause_ids: Optional[Sequence[str]] = None,
) -> Dict[str, Any]:
    state = _ensure_session(session_id)
    latest_assistant = None
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("clause_table"):
            latest_assistant = msg
            break
    if not latest_assistant:
        raise ValueError("No clause analysis found in session. Please run analysis first.")

    clause_table = latest_assistant.get("clause_table") or []
    if not clause_table:
        raise ValueError("No clause table available to export.")
    from . import config as agent_config

    all_edit_instructions = _build_edit_instructions_from_clause_table(clause_table)
    edit_instructions = _filter_edit_instructions_for_hitl(
        all_edit_instructions,
        accepted_clause_ids=accepted_clause_ids,
        require_hitl=agent_config.REQUIRE_REDLINE_HITL_ACCEPTANCE,
    )

    primary_artifact = _primary_readable_artifact(state)
    source_contract_text = (
        str(primary_artifact.extras.get("full_text") or "").strip()
        if primary_artifact is not None
        else "\n\n".join(
            [str(f.extras.get("full_text") or "") for f in state.files if f.extras.get("full_text")]
        ).strip()
    )
    source_docx_path = None
    if primary_artifact is not None and primary_artifact.path.suffix.lower() == ".docx" and primary_artifact.path.exists():
        source_docx_path = primary_artifact.path

    # When no DOCX found (PDF-only upload), try pdf2docx conversion before redlining
    if source_docx_path is None:
        pdf_candidates = [primary_artifact] if primary_artifact is not None else list(reversed(state.files))
        for artifact in pdf_candidates:
            if artifact is None:
                continue
            if artifact.path.suffix.lower() == ".pdf" and artifact.path.exists():
                try:
                    from pdf2docx import Converter
                    output_dir = UPLOAD_ROOT / session_id
                    converted_path = output_dir / f"{artifact.path.stem}_converted.docx"
                    cv = Converter(str(artifact.path))
                    cv.convert(str(converted_path))
                    cv.close()
                    source_docx_path = converted_path
                except Exception:
                    source_docx_path = None  # fallback to _build_doc_from_extracted_text
                break

    _validate_redline_template_source(primary_artifact, source_docx_path, session_id)
    from .master_orchestrator import log_redline_export_roster
    from .redline_docx import build_reviewed_contract_docx, build_reviewed_contract_docx_parallel

    log_redline_export_roster()
    output_path = UPLOAD_ROOT / session_id / "reviewed_contract.docx"
    render_warnings: List[str] = []
    verification_report: List[Dict[str, str]] = []
    if (
        agent_config.REQUIRE_REDLINE_HITL_ACCEPTANCE
        and len(all_edit_instructions) > 0
        and len(edit_instructions) == 0
    ):
        render_warnings.append(
            "Human review: no clauses selected for body redlines (or none matched); exported file "
            "matches the source contract with no strike/underline edits from this analysis."
        )
    builder = build_reviewed_contract_docx_parallel if agent_config.ENABLE_PARALLEL_REDLINE else build_reviewed_contract_docx
    builder(
        output_path=output_path,
        clause_table=clause_table,
        edit_instructions=edit_instructions,
        source_contract_text=source_contract_text,
        source_docx_path=source_docx_path,
        render_warnings=render_warnings,
        verification_flags=verification_report,
    )
    if render_warnings:
        latest_assistant["docx_export_warnings"] = render_warnings
        logger.warning("DOCX export validator warnings: %s", " | ".join(render_warnings))
    if verification_report:
        latest_assistant["verification_report"] = verification_report
        logger.info("Agent4 flagged %d edit(s) for review.", len(verification_report))
    return {
        "path": str(output_path),
        "filename": output_path.name,
        "warnings": render_warnings,
        "verification_report": verification_report,
    }


def generate_contract_commentary_docx(session_id: str) -> Dict[str, Any]:
    """Export primary contract DOCX with margin comments (no redlines). Requires analysis + DOCX/PDF primary."""
    from . import config as agent_config

    if not getattr(agent_config, "ENABLE_CONTRACT_COMMENTARY_DOCX", True):
        raise ValueError("Contract commentary DOCX export is disabled (ENABLE_CONTRACT_COMMENTARY_DOCX).")

    state = _ensure_session(session_id)
    latest_assistant = None
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("clause_table"):
            latest_assistant = msg
            break
    if not latest_assistant:
        raise ValueError("No clause analysis found in session. Please run analysis first.")

    clause_table = latest_assistant.get("clause_table") or []
    if not clause_table:
        raise ValueError("No clause table available to export.")

    primary_artifact = _primary_readable_artifact(state)
    primary_name = primary_artifact.name if primary_artifact is not None else ""
    source_contract_text = (
        str(primary_artifact.extras.get("full_text") or "").strip()
        if primary_artifact is not None
        else "\n\n".join(
            [str(f.extras.get("full_text") or "") for f in state.files if f.extras.get("full_text")]
        ).strip()
    )
    supporting_doc_texts: Dict[str, str] = {
        f.name: str(f.extras.get("full_text") or "")
        for f in state.files
        if f.name != primary_name and f.extras.get("full_text")
    }
    uploaded_filenames = [f.name for f in state.files]

    source_docx_path = None
    if primary_artifact is not None and primary_artifact.path.suffix.lower() == ".docx" and primary_artifact.path.exists():
        source_docx_path = primary_artifact.path
    if source_docx_path is None and primary_artifact is not None and primary_artifact.path.suffix.lower() == ".pdf":
        try:
            from pdf2docx import Converter

            output_dir = UPLOAD_ROOT / session_id
            converted_path = output_dir / f"{primary_artifact.path.stem}_converted.docx"
            if not converted_path.exists():
                cv = Converter(str(primary_artifact.path))
                cv.convert(str(converted_path))
                cv.close()
            source_docx_path = converted_path
        except Exception:
            source_docx_path = None

    from .redline_docx import build_contract_commentary_docx

    output_path = UPLOAD_ROOT / session_id / "contract_with_review_comments.docx"
    render_warnings: List[str] = []
    build_contract_commentary_docx(
        output_path=output_path,
        clause_table=clause_table,
        source_contract_text=source_contract_text,
        source_docx_path=source_docx_path,
        supporting_doc_texts=supporting_doc_texts or None,
        render_warnings=render_warnings,
        uploaded_filenames=uploaded_filenames,
        commentary_export_style=getattr(agent_config, "CONTRACT_COMMENTARY_EXPORT_STYLE", "counsel_bubble"),
        primary_upload_display_name=primary_name,
    )
    if render_warnings:
        logger.warning("Commentary DOCX warnings: %s", " | ".join(render_warnings))
    return {"path": str(output_path), "filename": output_path.name, "warnings": render_warnings}


def _build_edit_instructions_from_clause_table(clause_table: List[Dict[str, str]]) -> List[Dict[str, str]]:
    instructions: List[Dict[str, str]] = []
    for idx, row in enumerate(clause_table or [], start=1):
        evidence_text = (
            (row.get("evidence_quote") or row.get("evidence_snippet") or "").strip()
        )
        uploaded_position = (row.get("uploaded_position") or "").strip()
        original_text = evidence_text.split(" | ")[0].strip() if evidence_text else ""
        if uploaded_position and len(uploaded_position) > len(original_text):
            original_text = uploaded_position
        suggested_text = (row.get("suggested_text") or "").strip() or (row.get("gb_ideal_position") or "").strip()
        brief_cf = _brief_counterfactual_narrative(row)
        instructions.append(
            {
                "clause_id": str(idx),
                "clause_name": (row.get("clause_name") or "").strip(),
                "risk_level": (row.get("risk_level") or "Amber").strip(),
                "evidence_text": evidence_text,
                "original_text": original_text,
                "uploaded_position": uploaded_position,
                "suggested_text": suggested_text,
                "gb_ideal_position": (row.get("gb_ideal_position") or "").strip(),
                "mitigation_recommendation": (row.get("mitigation_recommendation") or "").strip(),
                "risk_rationale": (row.get("risk_rationale") or "").strip(),
                "reason": (row.get("risk_trigger") or row.get("risk_rationale") or "").strip(),
                "counterfactual": (row.get("counterfactual") or "").strip(),
                "brief_counterfactual": brief_cf,
                "approval_path": (row.get("approval_path") or "").strip(),
                "detected": (row.get("detected") or "Unclear").strip(),
                "confidence_score": str(row.get("confidence_score") or "").strip(),
            }
        )
    return instructions


def generate_clause_table_csv(session_id: str) -> Dict[str, Any]:
    state = _ensure_session(session_id)
    latest_assistant = None
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("clause_table"):
            latest_assistant = msg
            break
    if not latest_assistant:
        raise ValueError("No clause analysis found in session. Please run analysis first.")
    clause_table = latest_assistant.get("clause_table") or []
    if not clause_table:
        raise ValueError("No clause table available to export.")

    output_path = UPLOAD_ROOT / session_id / "clause_analysis.csv"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    headers = [
        "Clause",
        "Detected",
        "Uploaded Position",
        "GB Ideal Position",
        "Risk",
        "Confidence Label",
        "Confidence Score",
        "Risk Rationale",
        "Mitigation",
        "Approval Path",
        "Evidence",
        "Evidence Quote",
        "Evidence Source",
        "Evidence Location",
        "Anchoring Warning",
        "Knowledge Ref",
        "Counterfactual",
    ]
    with output_path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(headers)
        for row in clause_table:
            writer.writerow(
                [
                    row.get("clause_name", ""),
                    row.get("detected", ""),
                    row.get("uploaded_position", ""),
                    row.get("gb_ideal_position", ""),
                    row.get("risk_level", ""),
                    row.get("confidence_label", ""),
                    row.get("confidence_score", ""),
                    row.get("risk_rationale", ""),
                    row.get("mitigation_recommendation", ""),
                    row.get("approval_path", ""),
                    row.get("evidence_snippet", ""),
                    row.get("evidence_quote", ""),
                    row.get("evidence_source", ""),
                    row.get("evidence_location_hint", ""),
                    row.get("anchoring_warning", ""),
                    row.get("knowledge_reference", ""),
                    row.get("counterfactual", ""),
                ]
            )
    return {"path": str(output_path), "filename": output_path.name}


def generate_counterfactuals_csv(session_id: str) -> Dict[str, Any]:
    """Export counterfactual rows to CSV."""
    state = _ensure_session(session_id)
    latest_assistant = None
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("clause_table"):
            latest_assistant = msg
            break
    if not latest_assistant:
        raise ValueError("No clause analysis found in session. Please run analysis first.")
    clause_table = latest_assistant.get("clause_table") or []
    counterfactual_rows = latest_assistant.get("counterfactual_rows")
    if not counterfactual_rows and clause_table:
        counterfactual_rows = _counterfactual_rows_from_clause_table(clause_table)
    if not counterfactual_rows:
        raise ValueError("No counterfactuals available to export.")

    output_path = UPLOAD_ROOT / session_id / "counterfactuals.csv"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    headers = ["Clause", "Current Text", "What-if", "Expected Risk Shift", "Why"]
    with output_path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(headers)
        for row in counterfactual_rows:
            writer.writerow([
                row.get("clause", ""),
                row.get("current_text", ""),
                row.get("required_change", ""),
                row.get("risk_shift", ""),
                row.get("reason", ""),
            ])
    return {"path": str(output_path), "filename": "counterfactuals.csv"}


def generate_mitigation_checklist_csv(session_id: str) -> Dict[str, Any]:
    """Export mitigation checklist rows to CSV."""
    state = _ensure_session(session_id)
    latest_assistant = None
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("clause_table"):
            latest_assistant = msg
            break
    if not latest_assistant:
        raise ValueError("No clause analysis found in session. Please run analysis first.")
    checklist_rows = latest_assistant.get("mitigation_checklist_rows")
    if not checklist_rows:
        clause_table = latest_assistant.get("clause_table") or []
        checklist_rows = _checklist_rows_from_clause_table(clause_table)
    if not checklist_rows:
        raise ValueError("No mitigation checklist available to export.")

    output_path = UPLOAD_ROOT / session_id / "mitigation_checklist.csv"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    headers = ["Clause", "Risk Trigger", "Mitigation Recommendation", "Approval Route", "Priority"]
    with output_path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(headers)
        for row in checklist_rows:
            writer.writerow([
                row.get("clause", ""),
                row.get("risk_trigger") or row.get("risk", ""),
                row.get("mitigation", ""),
                row.get("approval_path", ""),
                row.get("priority", ""),
            ])
    return {"path": str(output_path), "filename": "mitigation_checklist.csv"}


def generate_verification_report_csv(session_id: str) -> Dict[str, Any]:
    """Export Agent 4 verification flags (flagged edits) to CSV. Available after Generate Reviewed DOCX."""
    state = _ensure_session(session_id)
    verification_report: List[Dict[str, str]] = []
    for msg in reversed(state.history):
        if msg.get("role") == "assistant" and msg.get("verification_report"):
            verification_report = msg["verification_report"]
            break
    if not verification_report:
        raise ValueError(
            "No verification report available. Generate Reviewed DOCX first. "
            "Agent 4 flags edits for semantic/contextual review when enabled."
        )

    output_path = UPLOAD_ROOT / session_id / "verification_report.csv"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    headers = ["Clause", "Reason", "Suggestion"]
    with output_path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(headers)
        for row in verification_report:
            writer.writerow([
                row.get("clause_name", ""),
                row.get("reason", ""),
                row.get("suggestion", ""),
            ])
    return {"path": str(output_path), "filename": "verification_report.csv"}


def export_session_state(session_id: str) -> Dict[str, Any]:
    """Serialize session state to persist across requests (e.g., in cache)."""
    state = _ensure_session(session_id)
    _load_chat_history_from_disk_if_empty(session_id, state)
    return {
        "files": [
            {
                "name": f.name,
                "path": str(f.path),
                "size": f.size,
                "mime": f.mime,
                "text_preview": f.text_preview,
                "image_bytes": [_encode_bytes(b) for b in f.extras.get("image_bytes", [])],
            }
            for f in state.files
        ],
        "history": list(state.history),
    }


def restore_session_state(session_id: str, payload: Dict[str, Any]) -> None:
    """Rehydrate state that was produced by `export_session_state`."""
    files: List[UploadedArtifact] = []
    for entry in payload.get("files", []):
        path = Path(entry.get("path", ""))
        files.append(
            UploadedArtifact(
                name=entry.get("name", path.name),
                path=path,
                size=entry.get("size"),
                mime=entry.get("mime"),
                text_preview=entry.get("text_preview"),
                extras={
                    "image_bytes": [
                        b for b in (_decode_bytes(x) for x in entry.get("image_bytes") or []) if b
                    ]
                },
            )
        )

    state = SessionState(
        files=files,
        history=list(payload.get("history", [])),
    )
    _SESSION_DATA[session_id] = state


def _encode_bytes(blob: Optional[bytes]) -> Optional[str]:
    if not blob:
        return None
    try:
        return base64.b64encode(blob).decode("ascii")
    except Exception:
        return None


def _decode_bytes(value: Optional[str]) -> Optional[bytes]:
    if not value:
        return None
    try:
        return base64.b64decode(value.encode("ascii"))
    except Exception:
        return None
