#!/usr/bin/env python3
"""
Automated accuracy test harness for the legal clause extraction pipeline.
Runs on all files in data/ and validates evidence-clause relevance.
Target: 90%+ accuracy (evidence validation pass rate).
"""

from __future__ import annotations

import json
import os
import sys
from pathlib import Path

# Add project root for imports
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(PROJECT_ROOT)

# Primary anchors: clause must have at least one of these in evidence (clause-specific)
_PRIMARY_ANCHORS: dict[str, list[str]] = {
    "limitation of liability": ["aggregate liability", "liability cap", "100% of", "200% of", "consequential", "indirect damages", "liability"],
    "governing law": ["governing law", "jurisdiction", "laws of"],
    "dispute resolution": ["arbitration", "arbitrator", "mediation", "siac", "lcia", "icc"],
    "firm price": ["firm", "fixed price"],
    "force majeure": ["force majeure", "impediment", "act of god"],
    "liquidated damages": ["liquidated damages", "ld ", "per week"],
    "orders extending": ["termination", "purchase orders", "in-effect"],
    "quantity protection": ["forecast", "quantity", "deviation", "+/-20%", "reimbursement"],
    "inventory requirements": ["inventory", "raw material", "weeks", "forecast", "finished goods"],
    "change orders": ["change order", "equitable adjustment", "nre", "fai"],
    "aerospace business critical": ["aerospace", "actuator", "schedule a", "technical specifications"],
}


def _primary_anchor_for_clause(name: str) -> list[str]:
    n = (name or "").lower()
    for key, terms in _PRIMARY_ANCHORS.items():
        if key in n:
            return terms
    return []


def validate_evidence(clause_name: str, evidence: str, detected: str) -> tuple[bool, str]:
    """
    Validate that evidence is relevant to the clause.
    Returns (pass: bool, reason: str).
    """
    if not evidence or detected != "Yes":
        # Unclear/No is acceptable - we're not claiming wrong attribution
        return True, "N/A (no evidence)"
    ev_lower = (evidence or "").lower()
    primary = _primary_anchor_for_clause(clause_name)
    if not primary:
        return True, "No primary anchor defined"
    has_anchor = any(p in ev_lower for p in primary if p)
    if not has_anchor:
        return False, f"Evidence lacks primary anchor for {clause_name}"
    return True, "OK"


def extract_text(filepath: Path, use_ocr_fallback: bool = True) -> str:
    """Extract text from PDF or DOCX. For scanned PDFs, optionally use OCR when standard extraction yields < 100 chars."""
    ext = filepath.suffix.lower()
    try:
        if ext == ".pdf":
            import fitz
            pages = []
            with fitz.open(str(filepath)) as doc:
                for i, page in enumerate(doc):
                    t = (page.get_text("text") or "").strip()
                    if t:
                        pages.append(f"[PAGE {i+1}]\n{t}")
            text = "\n\n".join(pages) if pages else ""
            # OCR fallback for scanned PDFs (requires Tesseract: brew install tesseract)
            use_ocr = use_ocr_fallback and len(text) < 100 and os.environ.get("USE_PDF_OCR", "0") == "1"
            if use_ocr and hasattr(fitz.Page, "get_textpage_ocr"):
                try:
                    ocr_pages = []
                    with fitz.open(str(filepath)) as doc:
                        for i, page in enumerate(doc):
                            if i >= 20:  # Limit to first 20 pages for speed
                                break
                            tp = page.get_textpage_ocr(language="eng")
                            t = (page.get_text("text", textpage=tp) or "").strip()
                            if t:
                                ocr_pages.append(f"[PAGE {i+1}]\n{t}")
                    if ocr_pages:
                        text = "\n\n".join(ocr_pages)
                except Exception:
                    pass
            return text
        elif ext in (".docx", ".docm"):
            from docx import Document as DocxDocument
            doc = DocxDocument(str(filepath))
            chunks = [p.text.strip() for p in doc.paragraphs if (p.text or "").strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tx = (cell.text or "").strip()
                        if tx:
                            chunks.append(tx)
            return "\n\n".join(chunks) if chunks else ""
        else:
            return ""
    except Exception as e:
        print(f"  WARNING: extract failed for {filepath.name}: {e}")
        return ""


def _process_one_file(
    fp: Path, knowledge_payload: dict
) -> tuple[str, dict]:
    """Process a single file; returns (filename, result_dict)."""
    import agent1_clause_analyzer as agent1

    text = extract_text(fp)
    if not text or len(text) < 100:
        return fp.name, {"error": "Empty or minimal text (possibly scanned PDF)", "rows": []}

    try:
        rows = agent1.run_aerospace_clause_extraction(
            contract_text=text,
            knowledge_payload=knowledge_payload,
        )
    except Exception as e:
        return fp.name, {"error": str(e), "rows": []}

    validated = failed = unclear_ok = 0
    details = []
    for row in rows:
        clause = row.get("clause_name", "")
        evidence = (row.get("evidence_snippet") or row.get("uploaded_position") or "").strip()
        detected = (row.get("detected") or "Unclear").strip()
        conf = float(row.get("confidence_score") or 0)
        pass_val, reason = validate_evidence(clause, evidence, detected)
        if detected != "Yes":
            unclear_ok += 1
        elif pass_val:
            validated += 1
        else:
            failed += 1
        details.append({
            "clause": clause,
            "detected": detected,
            "confidence": conf,
            "validation": "PASS" if pass_val else "FAIL",
            "reason": reason,
            "evidence_preview": (evidence[:80] + "...") if len(evidence) > 80 else evidence,
        })

    return fp.name, {
        "rows": len(rows),
        "validated": validated,
        "failed": failed,
        "unclear_ok": unclear_ok,
        "details": details,
    }


def run_tests(data_dir: Path, knowledge_path: Path, parallel: bool = True) -> dict:
    """Run clause extraction on all files and validate. Uses parallel processing when parallel=True."""
    from agents.sample_agent.knowledge_loader import load_knowledge_payload
    from concurrent.futures import ThreadPoolExecutor, as_completed

    knowledge_payload = load_knowledge_payload(knowledge_path)
    if not knowledge_payload.get("clauses"):
        return {"error": "No knowledge clauses loaded", "files": {}}

    files = list(data_dir.glob("*.pdf")) + list(data_dir.glob("*.docx"))
    files = sorted([f for f in files if f.suffix.lower() in (".pdf", ".docx")])

    results: dict = {"files": {}, "summary": {"total_clauses": 0, "validated": 0, "failed": 0, "unclear_ok": 0}}

    if parallel and len(files) > 1:
        with ThreadPoolExecutor(max_workers=min(5, len(files))) as ex:
            futures = {ex.submit(_process_one_file, fp, knowledge_payload): fp for fp in files}
            for fut in as_completed(futures):
                fname, finfo = fut.result()
                results["files"][fname] = finfo
    else:
        for fp in files:
            fname, finfo = _process_one_file(fp, knowledge_payload)
            results["files"][fname] = finfo

    for finfo in results["files"].values():
        if "error" in finfo:
            continue
        results["summary"]["total_clauses"] += finfo.get("rows", 0)
        results["summary"]["validated"] += finfo.get("validated", 0)
        results["summary"]["failed"] += finfo.get("failed", 0)
        results["summary"]["unclear_ok"] += finfo.get("unclear_ok", 0)

    total = results["summary"]["total_clauses"]
    yes_count = results["summary"]["validated"] + results["summary"]["failed"]
    if yes_count > 0:
        results["summary"]["accuracy_pct"] = round(
            100.0 * results["summary"]["validated"] / yes_count, 1
        )
    else:
        results["summary"]["accuracy_pct"] = 0.0
    # Overall: (validated + unclear) / total = pass rate
    if total > 0:
        results["summary"]["overall_pass_pct"] = round(
            100.0 * (results["summary"]["validated"] + results["summary"]["unclear_ok"]) / total, 1
        )
    else:
        results["summary"]["overall_pass_pct"] = 0.0

    return results


def main():
    data_dir = PROJECT_ROOT / "data"
    try:
        from agents.sample_agent import config as agent_config
        knowledge_path = Path(agent_config.POC_KNOWLEDGE_PATH)
    except Exception:
        knowledge_path = PROJECT_ROOT / "docs" / "knowledge" / "Contract_Positions_POC_GB_Legal_2026-02-19.json"
    if not knowledge_path.exists():
        knowledge_path = PROJECT_ROOT / "docs" / "knowledge" / "[RFP Tenders-LTAs] Contract Positions for POC_GB Legal_2026.02.19-v1.docx"
    if not knowledge_path.exists():
        knowledge_path = PROJECT_ROOT / "docs" / "knowledge" / "Contract_Positions_POC_GB_Legal_2026-02-19.json"

    if not data_dir.exists():
        print(f"Data dir not found: {data_dir}")
        sys.exit(1)
    if not knowledge_path.exists():
        print(f"Knowledge not found: {knowledge_path}")
        sys.exit(1)

    print("Running accuracy tests...")
    results = run_tests(data_dir, knowledge_path)

    if "error" in results:
        print(f"ERROR: {results['error']}")
        sys.exit(1)

    summary = results["summary"]
    print("\n" + "=" * 60)
    print("ACCURACY REPORT")
    print("=" * 60)
    print(f"Total clauses evaluated: {summary['total_clauses']}")
    print(f"Validated (Yes + pass):  {summary['validated']}")
    print(f"Failed (Yes + fail):     {summary['failed']}")
    print(f"Unclear/No (OK):        {summary['unclear_ok']}")
    print(f"Accuracy (of Yes):      {summary['accuracy_pct']}%")
    print(f"Overall pass rate:      {summary['overall_pass_pct']}%")
    print("=" * 60)

    for fname, finfo in results["files"].items():
        if "error" in finfo:
            print(f"\n{fname}: ERROR - {finfo['error']}")
        else:
            print(f"\n{fname}: validated={finfo['validated']}, failed={finfo['failed']}, unclear={finfo['unclear_ok']}")
            for d in finfo.get("details", []):
                if d["validation"] == "FAIL":
                    print(f"  FAIL: {d['clause'][:40]}... | {d['reason']}")

    out_dir = PROJECT_ROOT / "test_results"
    out_dir.mkdir(exist_ok=True)
    from datetime import datetime
    out_file = out_dir / f"accuracy_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(out_file, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"\nReport saved to: {out_file}")

    if summary["accuracy_pct"] >= 90:
        print("\n*** TARGET MET: 90%+ accuracy ***")
    else:
        print(f"\n*** Below target ({summary['accuracy_pct']}% < 90%). Review failures and iterate. ***")
        sys.exit(1)


if __name__ == "__main__":
    main()
