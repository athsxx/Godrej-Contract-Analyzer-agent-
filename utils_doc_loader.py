# utils_doc_loader.py

import fitz  # PyMuPDF
import docx2txt
import os
import logging
import shutil
import subprocess
from docx import Document as DocxDocument  # type: ignore

try:
    import textract  # type: ignore
except Exception:  # pragma: no cover - environment specific
    textract = None


logger = logging.getLogger(__name__)

DOCX_LIKE_EXTS = {".docx", ".docm", ".dotx", ".dotm"}
DOC_EXTS = {".doc", ".dot", ".rtf"}


def extract_text(filepath, *, include_page_markers: bool = False):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_pdf_text(filepath, include_page_markers=include_page_markers)
    elif ext in DOCX_LIKE_EXTS | DOC_EXTS:
        return extract_doc_text(filepath, include_page_markers=include_page_markers)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _apply_page_markers(pages_text: list[str], *, include_page_markers: bool) -> str:
    if not pages_text:
        return ""
    if not include_page_markers:
        return "\n".join(t for t in pages_text if t).strip()
    marked: list[str] = []
    for idx, text in enumerate(pages_text, start=1):
        if not text:
            continue
        marked.append(f"[PAGE {idx}]\n{text}")
    return "\n".join(marked).strip()


def extract_pdf_text(filepath, *, include_page_markers: bool = False):
    pages_text: list[str] = []
    try:
        with fitz.open(filepath) as doc:
            for page in doc:
                try:
                    pages_text.append(page.get_text("text") or "")
                except Exception as page_err:
                    logger.warning("Skipping unreadable PDF page: %s", page_err)
    except Exception as primary_err:
        logger.warning("Primary PDF read failed, attempting PyPDF2 fallback: %s", primary_err)
        try:
            import PyPDF2  # type: ignore

            with open(filepath, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                for page in reader.pages:
                    try:
                        pages_text.append(page.extract_text() or "")
                    except Exception as page_err:
                        logger.warning("Skipping unreadable PDF page (fallback): %s", page_err)
        except Exception as fallback_err:
            raise RuntimeError(f"PDF text extraction failed: {fallback_err}") from fallback_err

    cleaned = _apply_page_markers(pages_text, include_page_markers=include_page_markers)
    if not cleaned:
        raise RuntimeError("PDF text extraction returned empty content.")
    return cleaned

def render_pdf_images(filepath: str, *, max_pages: int = 3, dpi: int = 150) -> list[tuple[int, bytes]]:
    if max_pages <= 0:
        return []
    images: list[tuple[int, bytes]] = []
    try:
        with fitz.open(filepath) as doc:
            for idx, page in enumerate(doc, start=1):
                if len(images) >= max_pages:
                    break
                try:
                    pix = page.get_pixmap(dpi=dpi)
                    images.append((idx, pix.tobytes("png")))
                except Exception as page_err:
                    logger.warning("Skipping PDF page image render: %s", page_err)
    except Exception as exc:
        logger.warning("PDF image render failed for %s: %s", filepath, exc)
        return []
    return images

def extract_docx_text(filepath):
    return docx2txt.process(filepath)


def _extract_doc_with_textract(filepath: str) -> str:
    if textract is None:
        return ""
    try:
        data = textract.process(filepath)
        if isinstance(data, (bytes, bytearray)):
            return data.decode("utf-8", errors="ignore")
        return str(data or "")
    except Exception as exc:
        logger.warning("textract failed for %s: %s", filepath, exc)
        return ""


def _extract_doc_with_cli(filepath: str) -> str:
    for cmd in ("antiword", "catdoc"):
        exe = shutil.which(cmd)
        if not exe:
            continue
        try:
            proc = subprocess.run(
                [exe, filepath],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=30,
                check=False,
            )
        except Exception as exc:
            logger.warning("%s failed for %s: %s", cmd, filepath, exc)
            continue
        output = (proc.stdout or "").strip()
        if output:
            return output
        if proc.stderr:
            logger.warning("%s returned no text for %s: %s", cmd, filepath, proc.stderr.strip())
    return ""


def _apply_doc_page_markers(text: str, *, include_page_markers: bool) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return ""
    if not include_page_markers:
        return cleaned
    if "\f" in cleaned:
        pages = [p.strip() for p in cleaned.split("\f") if p.strip()]
        if not pages:
            return ""
        marked = [f"[PAGE {idx}]\n{page}" for idx, page in enumerate(pages, start=1)]
        return "\n".join(marked).strip()
    return f"[PAGE 1]\n{cleaned}".strip()


def extract_doc_text(filepath: str, *, include_page_markers: bool = False) -> str:
    text = ""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in DOCX_LIKE_EXTS:
        try:
            doc = DocxDocument(filepath)
            chunks: list[str] = []
            chunks.extend([p.text.strip() for p in doc.paragraphs if (p.text or "").strip()])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tx = (cell.text or "").strip()
                        if tx:
                            chunks.append(tx)
            text = "\n".join(chunks).strip()
        except Exception as exc:
            logger.warning("python-docx extraction failed for %s: %s", filepath, exc)
            text = ""
    if ext in DOCX_LIKE_EXTS and not text:
        try:
            text = docx2txt.process(filepath)
        except Exception as exc:
            logger.warning("docx extraction failed for %s: %s", filepath, exc)
            text = ""
    if not text:
        text = _extract_doc_with_textract(filepath)
        if not text:
            text = _extract_doc_with_cli(filepath)

    cleaned = _apply_doc_page_markers(text, include_page_markers=include_page_markers)
    if cleaned:
        return cleaned
    raise RuntimeError("DOC text extraction failed; please confirm the file is not corrupted or convert it to DOCX.")
